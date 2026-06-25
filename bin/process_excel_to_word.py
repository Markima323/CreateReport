#!/usr/bin/env python3
"""
Generate one Word value-analysis report per complete image-extracted JSON record.

Inputs default to:
1) bin/template/价值分析报告-自动生成基底模板.docx
2) bin/template/价值分析报告自动生成-Prompt.md
3) bin/template/价值分析报告生成规则.json
4) bin/json/图片提取数据.json
"""

from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import sys
import tempfile
import time
import urllib.error
import urllib.request
from copy import deepcopy
from dataclasses import dataclass
from datetime import date, datetime
from decimal import Decimal, InvalidOperation, ROUND_HALF_UP
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple
from zipfile import ZIP_DEFLATED, ZipFile
from xml.sax.saxutils import escape as xml_escape

from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

for stream in (sys.stdout, sys.stderr):
    if hasattr(stream, "reconfigure"):
        stream.reconfigure(encoding="utf-8", errors="replace")

try:
    from docx import Document
except Exception:  # pragma: no cover - import guard
    print(
        "Missing dependency: python-docx. Install with: pip install -r bin/requirements.txt",
        file=sys.stderr,
    )
    raise


DEFAULT_PROJECT_VALUES: Dict[str, Any] = {
    "项目年": 2025,
    "报告年": 2025,
    "报告批次": "1017",
    "分析基准日": "2025年6月30日",
    "报告日期": "2025年10月16日",
    "报告日期大写": "二〇二五年十月十六日",
    "有效期截止日": "2025年12月30日",
    "折现率": 0.03,
}

DEFAULT_GEMINI_MODEL = "gemini-3.5-flash"
MONEY_TOLERANCE = Decimal("1")
COEFFICIENT_TOLERANCE = Decimal("0.0001")
PERCENT_RATIO_TOLERANCE = Decimal("0.0001")

FIELD_COLUMNS: Dict[str, int] = {
    "序号": 1,
    "支行": 2,
    "姓名": 3,
    "基准日本金": 4,
    "基准日利息": 5,
    "基准日本息": 6,
    "地址": 7,
    "面积": 8,
    "第一顺位": 9,
    "登记类型": 10,
    "性质": 11,
    "权利价值": 12,
    "预计回收": 13,
    "评估单价": 14,
    "评估价值": 15,
    "快速变现系数": 16,
    "快速变现价值": 17,
    "清收完成时间": 18,
    "折现期限": 19,
    "折现系数": 20,
    "折现价值": 21,
    "处置费用": 22,
    "安置费用": 23,
    "可回收金额": 24,
    "债权评估值": 25,
    "偿债率": 26,
    "权证编号": 29,
    "证载权利人": 30,
    "保证人": 51,
}

OPTIONAL_GUARANTOR_COLUMNS: Dict[str, str] = {
    "保证人统一社会信用代码": "统一社会信用代码",
    "保证人类型": "类型",
    "保证人法定代表人": "法定代表人",
    "保证人成立日期": "成立日期",
    "保证人营业场所": "营业场所",
    "保证人经营范围": "经营范围",
}

DEBTOR_COLUMNS: List[Tuple[str, int, int, int, int, int]] = [
    ("债务人1", 31, 32, 33, 34, 35),
    ("债务人2", 36, 37, 38, 39, 40),
    ("债务人3", 41, 42, 43, 44, 45),
    ("债务人4", 46, 47, 48, 49, 50),
]

ILLEGAL_FILENAME_CHARS = re.compile(r'[\\/:\*\?"<>\|]')


@dataclass
class Debtor:
    source_index: int
    label: str
    name: str
    gender: str
    ethnicity: str
    id_number: str
    address: str


@dataclass
class PreparedReport:
    row_number: int
    sequence: str
    branch: str
    name: str
    report_number: str
    filename: str
    debtor_count: int
    guarantor_included: bool
    placement_cost_included: bool
    minimum_item: str
    replacements: Dict[str, str]
    conclusion_text: str
    guarantor_extra_lines: List[str]
    warnings: List[str]


@dataclass
class GeminiAudit:
    interaction_id: str
    model: str
    warnings: List[str]
    guarantor_details: Dict[str, Any]
    sources: List[str]
    usage: Dict[str, Any]


def is_blank(value: Any) -> bool:
    return value is None or str(value).strip() == ""


def clean_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return format_chinese_date(value)
    if isinstance(value, date):
        return format_chinese_date(value)
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).strip()


def format_chinese_date(value: date) -> str:
    return f"{value.year}年{value.month}月{value.day}日"


def to_decimal(value: Any, field_name: str) -> Decimal:
    if is_blank(value):
        raise ValueError(f"{field_name}为空")
    if isinstance(value, Decimal):
        return value
    if isinstance(value, (int, float)):
        return Decimal(str(value))
    text = str(value).strip().replace(",", "")
    if text.endswith("%"):
        return Decimal(text[:-1]) / Decimal("100")
    try:
        return Decimal(text)
    except InvalidOperation:
        raise ValueError(f"{field_name}不是有效数字: {value}")


def optional_decimal(value: Any, default: Decimal = Decimal("0")) -> Decimal:
    if is_blank(value):
        return default
    return to_decimal(value, "数值")


def round_half_up(value: Decimal, places: int = 0) -> Decimal:
    quant = Decimal("1") if places == 0 else Decimal("1").scaleb(-places)
    return value.quantize(quant, rounding=ROUND_HALF_UP)


def fmt_yuan(value: Decimal) -> str:
    return format(round_half_up(value, 2), ",.2f")


def fmt_wanyuan(value: Decimal) -> str:
    return format(round_half_up(value / Decimal("10000"), 2), ",.2f")


def fmt_area(value: Decimal) -> str:
    return format(round_half_up(value, 2), ",.2f")


def fmt_decimal(value: Decimal, places: int) -> str:
    return format(round_half_up(value, places), f",.{places}f")


def fmt_plain_number(value: Decimal) -> str:
    quantized = value.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    text = format(quantized.normalize(), "f")
    return text.rstrip("0").rstrip(".") if "." in text else text


def fmt_percent(value: Decimal) -> str:
    return f"{fmt_decimal(value * Decimal('100'), 2)}%"


def fmt_percent_compact(value: Decimal) -> str:
    return f"{fmt_plain_number(value * Decimal('100'))}%"


def sanitize_filename_part(value: str) -> Tuple[str, Optional[str]]:
    sanitized = ILLEGAL_FILENAME_CHARS.sub("-", value).strip()
    if sanitized != value:
        return sanitized, f"文件名非法字符已替换: {value} -> {sanitized}"
    return sanitized, None


def normalize_certificate_number(value: Any) -> str:
    return clean_text(value).replace("(", "（").replace(")", "）")


def id_card_checksum(first17: str) -> str:
    weights = [7, 9, 10, 5, 8, 4, 2, 1, 6, 3, 7, 9, 10, 5, 8, 4, 2]
    checks = "10X98765432"
    total = sum(int(ch) * weight for ch, weight in zip(first17, weights))
    return checks[total % 11]


def validate_or_repair_id_card(value: Any) -> Tuple[str, Optional[str]]:
    text = clean_text(value).replace(" ", "").upper()
    if re.fullmatch(r"\d{17}", text):
        repaired = text + id_card_checksum(text)
        return repaired, f"身份证号码17位已按校验位补齐: {text} -> {repaired}"
    if not re.fullmatch(r"\d{17}[\dX]", text):
        raise ValueError(f"身份证号码格式错误: {text}")
    birth = text[6:14]
    try:
        datetime.strptime(birth, "%Y%m%d")
    except ValueError:
        raise ValueError(f"身份证号码出生日期无效: {text}")
    expected = id_card_checksum(text[:17])
    if text[-1] != expected:
        raise ValueError(f"身份证号码校验位错误: {text}，应为 {expected}")
    return text, None


def normalize_ethnicity(value: Any) -> str:
    text = clean_text(value)
    if text and not text.endswith("族"):
        return text + "族"
    return text


def extract_floor(address: str) -> str:
    marker = address.rfind("室")
    if marker >= 0:
        before_room = address[:marker].strip()
        match = re.search(r"(\d+(?:[-－—]\d+)?)\s*$", before_room)
        if not match:
            raise ValueError(f"地址无法提取楼层，室号不明确: {address}")
        room = re.split(r"[-－—]", match.group(1))[-1]
    else:
        match = re.search(
            r"(?:"
            r"[-－—]\s*"
            r"|(?:幢|栋|座|号楼|单元)\s*"
            r")(\d{3,5})\s*$",
            address,
        )
        if not match:
            raise ValueError(
                "地址无法提取楼层，末尾没有明确的室号或“楼栋-房号”结构: "
                f"{address}"
            )
        room = match.group(1)
    if len(room) < 3:
        raise ValueError(f"地址无法提取楼层，房号位数不足: {address}")
    floor = int(room) // 100
    if floor <= 0:
        raise ValueError(f"地址无法提取楼层，房号异常: {address}")
    return f"{floor}层"


def choose_cached_or_computed(
    values: Dict[str, Any],
    field_name: str,
    computed: Decimal,
    tolerance: Decimal,
    warnings: List[str],
) -> Decimal:
    raw = values.get(field_name)
    if is_blank(raw):
        warnings.append(f"{field_name}为空，已由规则补算为 {computed}")
        return computed
    cached = to_decimal(raw, field_name)
    if abs(cached - computed) > tolerance:
        warnings.append(f"{field_name}与规则复核值不一致: 输入={cached}, 复核={computed}")
    return cached


def normalize_header(value: Any) -> str:
    return re.sub(r"\s+", "", clean_text(value))


def build_header_columns(ws: Any) -> Dict[str, int]:
    columns: Dict[str, int] = {}
    for col in range(1, ws.max_column + 1):
        header = normalize_header(ws.cell(1, col).value)
        if header:
            columns[header] = col
    return columns


def get_row_values(
    ws: Any,
    row_number: int,
    header_columns: Optional[Dict[str, int]] = None,
) -> Dict[str, Any]:
    values: Dict[str, Any] = {}
    for field, col in FIELD_COLUMNS.items():
        values[field] = ws.cell(row_number, col).value
    for i, (_, name_col, gender_col, ethnicity_col, id_col, addr_col) in enumerate(DEBTOR_COLUMNS, start=1):
        values[f"债务人{i}"] = ws.cell(row_number, name_col).value
        values[f"债务人{i}性别"] = ws.cell(row_number, gender_col).value
        values[f"债务人{i}民族"] = ws.cell(row_number, ethnicity_col).value
        values[f"债务人{i}身份证号码"] = ws.cell(row_number, id_col).value
        values[f"债务人{i}住址"] = ws.cell(row_number, addr_col).value
    for header in OPTIONAL_GUARANTOR_COLUMNS:
        col = (header_columns or {}).get(normalize_header(header))
        values[header] = ws.cell(row_number, col).value if col else None
    return values


def build_debtors(values: Dict[str, Any], warnings: List[str]) -> List[Debtor]:
    raw_debtors: List[Tuple[int, str, str, str, str, str]] = []
    for i in range(1, 5):
        name = clean_text(values.get(f"债务人{i}"))
        if not name:
            continue
        gender = clean_text(values.get(f"债务人{i}性别"))
        ethnicity = normalize_ethnicity(values.get(f"债务人{i}民族"))
        address = clean_text(values.get(f"债务人{i}住址"))
        missing = [
            label
            for label, item in [
                ("性别", gender),
                ("民族", ethnicity),
                ("身份证号码", values.get(f"债务人{i}身份证号码")),
                ("住址", address),
            ]
            if is_blank(item)
        ]
        if missing:
            raise ValueError(f"债务人{i}{name}缺少字段: {', '.join(missing)}")
        id_number, warning = validate_or_repair_id_card(values.get(f"债务人{i}身份证号码"))
        if warning:
            warnings.append(f"债务人{i}{name}: {warning}")
        raw_debtors.append((i, name, gender, ethnicity, id_number, address))

    if not raw_debtors:
        raise ValueError("未找到债务人")

    labels = ["债务人"] if len(raw_debtors) == 1 else ["债务人一", "债务人二", "债务人三", "债务人四"]
    debtors: List[Debtor] = []
    for index, item in enumerate(raw_debtors):
        source_index, name, gender, ethnicity, id_number, address = item
        debtors.append(
            Debtor(
                source_index=source_index,
                label=labels[index],
                name=name,
                gender=gender,
                ethnicity=ethnicity,
                id_number=id_number,
                address=address,
            )
        )
    return debtors


def build_prepared_report(
    row_number: int,
    values: Dict[str, Any],
    project: Dict[str, Any],
    guarantor_details: Optional[Dict[str, Any]] = None,
) -> PreparedReport:
    warnings: List[str] = []
    required = ["序号", "支行", "姓名", "地址", "面积", "权利价值", "预计回收", "评估单价", "权证编号", "证载权利人"]
    missing = [field for field in required if is_blank(values.get(field))]
    if missing:
        raise ValueError(f"必要字段为空: {', '.join(missing)}")

    sequence = clean_text(values["序号"])
    branch = clean_text(values["支行"])
    name = clean_text(values["姓名"])
    address = clean_text(values["地址"])
    floor = extract_floor(address)
    debtors = build_debtors(values, warnings)

    principal = to_decimal(values["基准日本金"], "基准日本金")
    interest = to_decimal(values["基准日利息"], "基准日利息")
    basis = choose_cached_or_computed(
        values,
        "基准日本息",
        principal + interest,
        MONEY_TOLERANCE,
        warnings,
    )
    if basis == 0:
        raise ValueError("基准日本息为0，无法计算偿债率")

    area = to_decimal(values["面积"], "面积")
    unit_price = to_decimal(values["评估单价"], "评估单价")
    market_value = choose_cached_or_computed(
        values,
        "评估价值",
        round_half_up(area * unit_price, 0),
        MONEY_TOLERANCE,
        warnings,
    )
    quick_coeff = to_decimal(values["快速变现系数"], "快速变现系数")
    quick_value = choose_cached_or_computed(
        values,
        "快速变现价值",
        round_half_up(market_value * quick_coeff, 0),
        MONEY_TOLERANCE,
        warnings,
    )
    discount_term = to_decimal(values["折现期限"], "折现期限")
    discount_rate = to_decimal(project["折现率"], "折现率")
    discount_coeff_calc = round_half_up(
        Decimal(str(float((Decimal("1") + discount_rate)) ** float(-discount_term))),
        4,
    )
    discount_coeff = choose_cached_or_computed(
        values,
        "折现系数",
        discount_coeff_calc,
        COEFFICIENT_TOLERANCE,
        warnings,
    )
    discounted_value = choose_cached_or_computed(
        values,
        "折现价值",
        round_half_up(quick_value * discount_coeff, 0),
        MONEY_TOLERANCE,
        warnings,
    )
    disposal_cost = optional_decimal(values.get("处置费用"))
    placement_cost = optional_decimal(values.get("安置费用"))
    recover_value = choose_cached_or_computed(
        values,
        "可回收金额",
        discounted_value - disposal_cost - placement_cost,
        MONEY_TOLERANCE,
        warnings,
    )
    right_value = to_decimal(values["权利价值"], "权利价值")
    debt_value_calc = min(basis, right_value, recover_value)
    debt_value = choose_cached_or_computed(
        values,
        "债权评估值",
        debt_value_calc,
        MONEY_TOLERANCE,
        warnings,
    )
    repayment_ratio = choose_cached_or_computed(
        values,
        "偿债率",
        debt_value / basis,
        PERCENT_RATIO_TOLERANCE,
        warnings,
    )

    property_nature = clean_text(values.get("性质"))
    has_placement_cost = placement_cost > 0
    if has_placement_cost:
        placement_title_suffix = "及安置费用"
        placement_sentence = (
            "同时，由于抵押物为债务人的唯一住房，需考虑安置费用，"
            f"根据委托人提供的资料，本次考虑安置费用为{fmt_yuan(placement_cost)}元。"
        )
        recover_formula = "抵押物折现价值-抵押物处置费用-安置费用"
        recover_calculation = f"{fmt_yuan(discounted_value)}-{fmt_yuan(disposal_cost)}-{fmt_yuan(placement_cost)}"
    else:
        placement_title_suffix = ""
        placement_sentence = "同时，由于抵押物并非债务人的唯一住房，无需考虑安置费用。"
        recover_formula = "抵押物折现价值-抵押物处置费用"
        recover_calculation = f"{fmt_yuan(discounted_value)}-{fmt_yuan(disposal_cost)}"

    if "非唯一" in property_nature and has_placement_cost:
        warnings.append("性质为非唯一住房，但安置费用大于0，已按安置费用生成")
    elif "唯一" in property_nature and not has_placement_cost:
        warnings.append("性质为唯一住房，但安置费用为空或0，已按安置费用生成")

    minimum_candidates = [
        ("抵押物可回收金额", recover_value),
        ("资产抵押债权金额", basis),
        ("抵押权利价值", right_value),
    ]
    minimum_value = min(item[1] for item in minimum_candidates)
    minimum_names = [label for label, amount in minimum_candidates if abs(amount - minimum_value) <= MONEY_TOLERANCE]
    if len(minimum_names) > 1:
        minimum_item = "与".join(minimum_names)
        warnings.append(f"{minimum_item}并列为三者最低值")
    else:
        minimum_item = minimum_names[0]

    conclusion = (
        "由于抵押物可偿债金额为抵押物可回收金额、资产抵押债权金额、抵押权利价值三者中孰低者，"
        f"抵押物可回收金额为{fmt_yuan(recover_value)}元，资产抵押债权金额为{fmt_yuan(basis)}元，"
        f"抵押权利价值为{fmt_yuan(right_value)}元，{minimum_item}"
    )
    if len(minimum_names) > 1:
        conclusion += f"并列为三者最低值，所以{name}户债权资产分析价值为{minimum_item}，即{fmt_yuan(debt_value)}元。"
    else:
        conclusion += f"为三者最低值，所以{name}户债权资产分析价值为{minimum_item}，即{fmt_yuan(debt_value)}元。"

    branch_for_filename, branch_warning = sanitize_filename_part(branch)
    name_for_filename, name_warning = sanitize_filename_part(name)
    for warning in [branch_warning, name_warning]:
        if warning:
            warnings.append(warning)

    report_year = clean_text(project["报告年"])
    report_batch = clean_text(project["报告批次"])
    report_number = clean_text(values.get("报告号"))
    if not report_number:
        report_number = (
            f"天地恒安[{report_year}]资评咨字第"
            f"{report_batch}-{sequence}号"
        )
    filename = f"价值分析报告-工行个贷不良资产-{branch_for_filename}-{name_for_filename}.docx"

    guarantor = clean_text(values.get("保证人"))
    debtor_names_overview = "及".join(debtor.name for debtor in debtors)
    if guarantor:
        debtor_names_overview = f"{debtor_names_overview}，债务责任关联方为{guarantor}"
    debtor_names_table = "、".join(debtor.name for debtor in debtors)

    guarantor_details = guarantor_details or {}
    guarantor_extra_lines = [
        clean_text(line) for line in guarantor_details.get("经营范围附加行", []) if clean_text(line)
    ]

    replacements: Dict[str, str] = {
        "保证人名称": guarantor,
        "保证人成立日期": clean_text(guarantor_details.get("成立日期")),
        "保证人法定代表人": clean_text(guarantor_details.get("法定代表人")),
        "保证人类型": clean_text(guarantor_details.get("类型")),
        "保证人经营范围": clean_text(guarantor_details.get("经营范围")),
        "保证人统一社会信用代码": clean_text(guarantor_details.get("统一社会信用代码")),
        "保证人营业场所": clean_text(guarantor_details.get("营业场所")),
        "债务人": debtor_names_table,
        "债务人姓名_及": debtor_names_overview,
        "债权评估值": fmt_yuan(debt_value),
        "偿债率百分比": fmt_percent(repayment_ratio),
        "分析基准日": clean_text(project["分析基准日"]),
        "利息万元": fmt_wanyuan(interest),
        "单价": fmt_yuan(unit_price),
        "可回收公式": recover_formula,
        "可回收计算式": recover_calculation,
        "可回收金额": fmt_yuan(recover_value),
        "坐落": address,
        "基准日本息": fmt_yuan(basis),
        "处置费用": fmt_yuan(disposal_cost),
        "姓名": name,
        "安置费用标题后缀": placement_title_suffix,
        "安置费用说明": placement_sentence,
        "层": floor,
        "市值": fmt_yuan(market_value),
        "序号": sequence,
        "快速变现价值": fmt_yuan(quick_value),
        "快速变现系数": fmt_plain_number(quick_coeff),
        "折现价值": fmt_yuan(discounted_value),
        "折现期限": fmt_plain_number(discount_term),
        "折现率百分比": fmt_percent_compact(discount_rate),
        "折现系数": fmt_decimal(discount_coeff, 4),
        "报告号": report_number,
        "报告年": report_year,
        "报告批次": report_batch,
        "报告日期": clean_text(project["报告日期"]),
        "报告日期大写": clean_text(project["报告日期大写"]),
        "最低值名称": minimum_item,
        "有效期截止日": clean_text(project["有效期截止日"]),
        "本息万元": fmt_wanyuan(basis),
        "本金万元": fmt_wanyuan(principal),
        "权利人": clean_text(values["证载权利人"]),
        "权利价值": fmt_yuan(right_value),
        "权利价值万元": fmt_wanyuan(right_value),
        "权证号": normalize_certificate_number(values["权证编号"]),
        "评估值万元": fmt_wanyuan(debt_value),
        "面积": fmt_area(area),
        "项目年": clean_text(project["项目年"]),
        "预计回收": clean_text(values["预计回收"]),
    }

    for i in range(1, 5):
        if i <= len(debtors):
            debtor = debtors[i - 1]
            replacements.update(
                {
                    f"债务人{i}": debtor.name,
                    f"债务人{i}住址": debtor.address,
                    f"债务人{i}性别": debtor.gender,
                    f"债务人{i}民族显示": debtor.ethnicity,
                    f"债务人{i}称谓": debtor.label,
                    f"债务人{i}身份证号码": debtor.id_number,
                }
            )
        else:
            replacements.update(
                {
                    f"债务人{i}": "",
                    f"债务人{i}住址": "",
                    f"债务人{i}性别": "",
                    f"债务人{i}民族显示": "",
                    f"债务人{i}称谓": "",
                    f"债务人{i}身份证号码": "",
                }
            )

    return PreparedReport(
        row_number=row_number,
        sequence=sequence,
        branch=branch,
        name=name,
        report_number=report_number,
        filename=filename,
        debtor_count=len(debtors),
        guarantor_included=bool(guarantor),
        placement_cost_included=has_placement_cost,
        minimum_item=minimum_item,
        replacements=replacements,
        conclusion_text=conclusion,
        guarantor_extra_lines=guarantor_extra_lines,
        warnings=warnings,
    )


def get_excel_guarantor_details(values: Dict[str, Any]) -> Dict[str, Any]:
    details: Dict[str, Any] = {}
    for header, detail_key in OPTIONAL_GUARANTOR_COLUMNS.items():
        value = clean_text(values.get(header))
        if value:
            details[detail_key] = value
    return details


def validate_guarantor_details(
    guarantor_name: str,
    details: Dict[str, Any],
) -> None:
    if not guarantor_name:
        return
    required_fields = (
        "统一社会信用代码",
        "类型",
        "法定代表人",
        "成立日期",
        "营业场所",
        "经营范围",
    )
    missing = [field for field in required_fields if not clean_text(details.get(field))]
    if missing:
        raise ValueError(
            f"保证人“{guarantor_name}”资料不完整，缺少: {', '.join(missing)}。"
            "请通过生成面板粘贴企业资料并使用 Gemini 提取保存。"
        )


def merge_gemini_guarantor_details(
    report: PreparedReport,
    excel_details: Dict[str, Any],
    audit: GeminiAudit,
) -> None:
    merged = dict(audit.guarantor_details)
    merged.update(excel_details)
    replacement_keys = {
        "统一社会信用代码": "保证人统一社会信用代码",
        "类型": "保证人类型",
        "法定代表人": "保证人法定代表人",
        "成立日期": "保证人成立日期",
        "营业场所": "保证人营业场所",
        "经营范围": "保证人经营范围",
    }
    for detail_key, replacement_key in replacement_keys.items():
        report.replacements[replacement_key] = clean_text(merged.get(detail_key))
    report.guarantor_extra_lines = [
        clean_text(line)
        for line in merged.get("经营范围附加行", [])
        if clean_text(line)
    ]
    report.warnings.extend(f"Gemini: {warning}" for warning in audit.warnings if clean_text(warning))


GEMINI_RESPONSE_SCHEMA: Dict[str, Any] = {
    "type": "object",
    "properties": {
        "warnings": {
            "type": "array",
            "items": {"type": "string"},
            "description": "Only concrete conflicts or risks found in the supplied generation plan.",
        },
        "guarantor": {
            "type": "object",
            "properties": {
                "name": {"type": "string"},
                "unified_social_credit_code": {"type": "string"},
                "company_type": {"type": "string"},
                "legal_representative": {"type": "string"},
                "establishment_date": {"type": "string"},
                "business_address": {"type": "string"},
                "business_scope": {"type": "string"},
                "business_scope_extra_lines": {
                    "type": "array",
                    "items": {"type": "string"},
                },
            },
            "required": [
                "name",
                "unified_social_credit_code",
                "company_type",
                "legal_representative",
                "establishment_date",
                "business_address",
                "business_scope",
                "business_scope_extra_lines",
            ],
        },
    },
    "required": ["warnings", "guarantor"],
}

GUARANTOR_EXTRACTION_SCHEMA: Dict[str, Any] = {
    "type": "object",
    "properties": {
        "guarantor": {"type": "string"},
        "unified_social_credit_code": {"type": "string"},
        "company_type": {"type": "string"},
        "legal_representative": {"type": "string"},
        "establishment_date": {"type": "string"},
        "business_address": {"type": "string"},
        "business_scope": {"type": "string"},
    },
    "required": [
        "guarantor",
        "unified_social_credit_code",
        "company_type",
        "legal_representative",
        "establishment_date",
        "business_address",
        "business_scope",
    ],
}


class GeminiInteractionsClient:
    endpoint = "https://generativelanguage.googleapis.com/v1beta/interactions"

    def __init__(self, api_key: str, timeout_seconds: int = 180) -> None:
        self.api_key = api_key
        self.timeout_seconds = timeout_seconds

    def create(self, **payload: Any) -> Dict[str, Any]:
        data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        for attempt in range(3):
            request = urllib.request.Request(
                self.endpoint,
                data=data,
                headers={
                    "Content-Type": "application/json; charset=utf-8",
                    "x-goog-api-key": self.api_key,
                },
                method="POST",
            )
            try:
                with urllib.request.urlopen(
                    request,
                    timeout=self.timeout_seconds,
                ) as response:
                    return json.loads(response.read().decode("utf-8"))
            except urllib.error.HTTPError as exc:
                body = exc.read().decode("utf-8", errors="replace")
                billing_failure = any(
                    marker in body.lower()
                    for marker in (
                        "prepayment credits are depleted",
                        "insufficient credits",
                        "billing account is not active",
                    )
                )
                if (
                    not billing_failure
                    and exc.code in {429, 500, 502, 503, 504}
                    and attempt < 2
                ):
                    time.sleep(2**attempt)
                    continue
                raise RuntimeError(f"Gemini API HTTP {exc.code}: {body}") from exc
            except urllib.error.URLError as exc:
                if attempt < 2:
                    time.sleep(2**attempt)
                    continue
                raise RuntimeError(f"无法连接 Gemini API: {exc.reason}") from exc
        raise RuntimeError("Gemini API 调用失败")


def create_gemini_client(api_key: str) -> GeminiInteractionsClient:
    return GeminiInteractionsClient(api_key=api_key)


def get_interaction_output_text(interaction: Dict[str, Any]) -> str:
    output_text = clean_text(interaction.get("output_text"))
    if output_text:
        return output_text
    text_blocks: List[str] = []
    for step in interaction.get("steps", []):
        if step.get("type") != "model_output":
            continue
        for content in step.get("content", []):
            if content.get("type") == "text" and content.get("text"):
                text_blocks.append(str(content["text"]))
    return "".join(text_blocks).strip()


def extract_guarantor_details_from_text(
    client: Any,
    model: str,
    thinking_level: str,
    guarantor_name: str,
    source_text: str,
    store_interactions: bool = False,
) -> GeminiAudit:
    guarantor_name = clean_text(guarantor_name)
    source_text = clean_text(source_text)
    if not guarantor_name:
        raise ValueError("保证人名称为空")
    if not source_text:
        raise ValueError("请先粘贴保证人资料")

    prompt = (
        "你是企业登记信息提取器。请只从用户粘贴的原始资料中提取字段，"
        "不得联网搜索，不得使用常识补全，不得改写公司名称。"
        "无法从原文确定的字段返回空字符串。经营范围应完整保留原文内容，"
        "删除网页菜单、广告、页脚和与目标公司无关的内容。\n\n"
        f"目标保证人：{guarantor_name}\n\n"
        "原始资料：\n"
        f"{source_text}"
    )
    interaction = client.create(
        model=model,
        input=prompt,
        response_format={
            "type": "text",
            "mime_type": "application/json",
            "schema": GUARANTOR_EXTRACTION_SCHEMA,
        },
        generation_config={"thinking_level": thinking_level},
        store=store_interactions,
    )
    if interaction.get("status") == "failed":
        raise RuntimeError(
            f"Gemini interaction 失败: {interaction.get('error') or interaction}"
        )
    output_text = get_interaction_output_text(interaction)
    if not output_text:
        raise RuntimeError("Gemini 返回了空响应")
    try:
        payload = json.loads(output_text)
    except json.JSONDecodeError as exc:
        raise RuntimeError(f"Gemini 返回的结构化 JSON 无法解析: {exc}") from exc

    returned_name = clean_text(payload.get("guarantor"))
    if returned_name != guarantor_name:
        raise ValueError(
            f"Gemini 提取的保证人名称不一致: 目标={guarantor_name}, "
            f"返回={returned_name or '空'}"
        )
    details = {
        "保证人": guarantor_name,
        "统一社会信用代码": clean_text(payload.get("unified_social_credit_code")),
        "类型": clean_text(payload.get("company_type")),
        "法定代表人": clean_text(payload.get("legal_representative")),
        "成立日期": clean_text(payload.get("establishment_date")),
        "营业场所": clean_text(payload.get("business_address")),
        "经营范围": clean_text(payload.get("business_scope")),
    }
    missing = [
        field
        for field in (
            "统一社会信用代码",
            "类型",
            "法定代表人",
            "成立日期",
            "营业场所",
            "经营范围",
        )
        if not details[field]
    ]
    warnings = [f"原始资料中未提取到: {', '.join(missing)}"] if missing else []
    usage = interaction.get("usage") if isinstance(interaction.get("usage"), dict) else {}
    return GeminiAudit(
        interaction_id=clean_text(interaction.get("id")),
        model=clean_text(interaction.get("model")) or model,
        warnings=warnings,
        guarantor_details=details,
        sources=[],
        usage=usage,
    )


def build_gemini_input(
    prompt_text: str,
    values: Dict[str, Any],
    report: PreparedReport,
    excel_details: Dict[str, Any],
) -> str:
    review_payload = {
        "document_type": "1",
        "row_number": report.row_number,
        "sequence": report.sequence,
        "calculation_plan": {
            "placement_cost_included": report.placement_cost_included,
            "minimum_item": report.minimum_item,
            "debtor_count": report.debtor_count,
            "guarantor_included": report.guarantor_included,
            "property_nature": clean_text(values.get("性质")),
            "expected_recovery_period": clean_text(values.get("预计回收")),
        },
        "guarantor_name": clean_text(values.get("保证人")),
        "excel_guarantor_details": excel_details,
    }
    return (
        f"{prompt_text}\n\n"
        "## 本次 Gemini 任务\n"
        "你只负责审核下面这条本地生成计划，并在保证人非空时核验 JSON 未提供的企业信息。"
        "不要输出报告正文，不要改写金额，不要推断债务人个人信息。"
        "本地 JSON 已提供的保证人字段具有最高优先级，返回时必须原样保留；"
        "未能从可靠网页核验的字段返回空字符串。"
        "如果保证人为空，guarantor 的所有字段均返回空值。\n\n"
        + json.dumps(review_payload, ensure_ascii=False, indent=2, default=str)
    )


def call_gemini(
    client: Any,
    model: str,
    thinking_level: str,
    prompt_text: str,
    values: Dict[str, Any],
    report: PreparedReport,
    excel_details: Dict[str, Any],
    use_google_search: bool,
    store_interactions: bool,
) -> GeminiAudit:
    kwargs: Dict[str, Any] = {
        "model": model,
        "input": build_gemini_input(prompt_text, values, report, excel_details),
        "response_format": {
            "type": "text",
            "mime_type": "application/json",
            "schema": GEMINI_RESPONSE_SCHEMA,
        },
        "generation_config": {"thinking_level": thinking_level},
        "store": store_interactions,
    }
    if report.guarantor_included and use_google_search:
        kwargs["tools"] = [{"type": "google_search"}]
    interaction = client.create(**kwargs)
    if interaction.get("status") == "failed":
        raise RuntimeError(
            f"Gemini interaction 失败: {interaction.get('error') or interaction}"
        )
    output_text = get_interaction_output_text(interaction)
    if not output_text:
        raise RuntimeError("Gemini 返回了空响应")
    try:
        payload = json.loads(output_text)
    except json.JSONDecodeError as exc:
        raise RuntimeError(f"Gemini 返回的结构化 JSON 无法解析: {exc}") from exc

    guarantor = payload.get("guarantor") or {}
    details: Dict[str, Any] = {
        "统一社会信用代码": clean_text(guarantor.get("unified_social_credit_code")),
        "类型": clean_text(guarantor.get("company_type")),
        "法定代表人": clean_text(guarantor.get("legal_representative")),
        "成立日期": clean_text(guarantor.get("establishment_date")),
        "营业场所": clean_text(guarantor.get("business_address")),
        "经营范围": clean_text(guarantor.get("business_scope")),
        "经营范围附加行": [
            clean_text(line)
            for line in guarantor.get("business_scope_extra_lines", [])
            if clean_text(line)
        ],
    }
    sources: List[str] = []
    for step in interaction.get("steps", []):
        if step.get("type") != "model_output":
            continue
        for content in step.get("content", []):
            if content.get("type") != "text":
                continue
            for annotation in content.get("annotations", []):
                if annotation.get("type") != "url_citation":
                    continue
                url = clean_text(annotation.get("url"))
                if url.lower().startswith(("https://", "http://")) and url not in sources:
                    sources.append(url)
    warnings = [
        clean_text(warning)
        for warning in payload.get("warnings", [])
        if clean_text(warning)
    ]
    expected_guarantor = clean_text(values.get("保证人"))
    returned_guarantor = clean_text(guarantor.get("name"))
    if expected_guarantor and returned_guarantor != expected_guarantor:
        warnings.append(
            f"保证人名称未精确匹配，已拒绝 Gemini 企业信息: "
            f"JSON={expected_guarantor}, Gemini={returned_guarantor or '空'}"
        )
        details = {}
        sources = []
    elif expected_guarantor and any(
        clean_text(value)
        for key, value in details.items()
        if key != "经营范围附加行"
    ) and not sources:
        warnings.append("Gemini 未提供可核验来源网址，已拒绝其保证人企业信息")
        details = {}
    usage = interaction.get("usage") if isinstance(interaction.get("usage"), dict) else {}
    return GeminiAudit(
        interaction_id=clean_text(interaction.get("id")),
        model=clean_text(interaction.get("model")) or model,
        warnings=warnings,
        guarantor_details=details,
        sources=sources,
        usage=usage,
    )


def remove_paragraph(paragraph: Any) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def set_paragraph_text(paragraph: Any, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        paragraph.runs[0].font.highlight_color = None
        for run in paragraph.runs[1:]:
            run.text = ""
            run.font.highlight_color = None
    else:
        paragraph.add_run(text)


def insert_paragraph_after(paragraph: Any, text: str) -> Any:
    new_element = OxmlElement("w:p")
    if paragraph._p.pPr is not None:
        new_element.append(deepcopy(paragraph._p.pPr))
    paragraph._p.addnext(new_element)
    new_paragraph = Paragraph(new_element, paragraph._parent)
    run = new_paragraph.add_run(text)
    if paragraph.runs and paragraph.runs[0]._r.rPr is not None:
        run._r.insert(0, deepcopy(paragraph.runs[0]._r.rPr))
    return new_paragraph


def iter_paragraphs(container: Any) -> Iterable[Any]:
    for paragraph in container.paragraphs:
        yield paragraph
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def iter_all_paragraphs(doc: Any) -> Iterable[Any]:
    yield from iter_paragraphs(doc)
    for section in doc.sections:
        for container in [
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ]:
            yield from iter_paragraphs(container)


def clear_highlights(doc: Any) -> None:
    for paragraph in iter_all_paragraphs(doc):
        for run in paragraph.runs:
            run.font.highlight_color = None


def apply_literal_adjustments(doc: Any, has_guarantor: bool) -> None:
    if has_guarantor:
        return
    for paragraph in iter_all_paragraphs(doc):
        text = paragraph.text
        if "保证人的偿债能力" not in text and "保证合同" not in text:
            continue
        adjusted = text.replace("、保证人的偿债能力", "")
        adjusted = adjusted.replace("贷款合同、借款凭据、抵押合同、保证合同等资料", "贷款合同、借款凭据、抵押合同等资料")
        adjusted = adjusted.replace("等五个方面", "等四个方面")
        if adjusted != text:
            set_paragraph_text(paragraph, adjusted)


def delete_optional_blocks(doc: Any, report: PreparedReport) -> None:
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        for debtor_index in range(report.debtor_count + 1, 5):
            if f"{{{{债务人{debtor_index}称谓}}}}" in text:
                remove_paragraph(paragraph)
                break

    if report.guarantor_included:
        optional_guarantor_fields = [
            ("统一社会信用代码：{{保证人统一社会信用代码}}", "保证人统一社会信用代码"),
            ("类    型：{{保证人类型}}", "保证人类型"),
            ("法定代表人：{{保证人法定代表人}}", "保证人法定代表人"),
            ("成立日期：{{保证人成立日期}}", "保证人成立日期"),
            ("营业场所：{{保证人营业场所}}", "保证人营业场所"),
            ("经营范围：{{保证人经营范围}}", "保证人经营范围"),
        ]
        removable_prefixes = [
            prefix for prefix, key in optional_guarantor_fields if not report.replacements.get(key)
        ]
    else:
        removable_prefixes = [
            "债务责任关联方简介",
            "保证人：{{保证人名称}}",
            "统一社会信用代码：{{保证人统一社会信用代码}}",
            "类    型：{{保证人类型}}",
            "法定代表人：{{保证人法定代表人}}",
            "成立日期：{{保证人成立日期}}",
            "营业场所：{{保证人营业场所}}",
            "经营范围：{{保证人经营范围}}",
        ]

    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if any(text.startswith(prefix) for prefix in removable_prefixes):
            remove_paragraph(paragraph)


def replace_placeholders_in_paragraph(paragraph: Any, replacements: Dict[str, str]) -> None:
    original = paragraph.text
    for run in paragraph.runs:
        run.font.highlight_color = None
        text = run.text
        if "{{" not in text:
            continue
        for key, value in replacements.items():
            text = text.replace("{{" + key + "}}", value)
        run.text = text

    if "{{" in paragraph.text:
        replaced = original
        for key, value in replacements.items():
            replaced = replaced.replace("{{" + key + "}}", value)
        if replaced != original:
            set_paragraph_text(paragraph, replaced)


def replace_placeholders(doc: Any, report: PreparedReport) -> None:
    for paragraph in iter_all_paragraphs(doc):
        if "由于抵押物可偿债金额为抵押物可回收金额" in paragraph.text:
            set_paragraph_text(paragraph, report.conclusion_text)
            continue
        replace_placeholders_in_paragraph(paragraph, report.replacements)


def add_guarantor_extra_lines(doc: Any, report: PreparedReport) -> None:
    if not report.guarantor_extra_lines:
        return
    in_guarantor_block = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("保证人："):
            in_guarantor_block = True
            continue
        if in_guarantor_block and text.startswith("委托合同约定的价值分析报告使用人"):
            return
        if in_guarantor_block and text.startswith("经营范围："):
            current = paragraph
            for line in report.guarantor_extra_lines:
                current = insert_paragraph_after(current, line)
            return


def postprocess_docx_xml(path: Path, replacements: Dict[str, str]) -> None:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as fp:
        temp_path = Path(fp.name)
    try:
        with ZipFile(path, "r") as source, ZipFile(temp_path, "w", ZIP_DEFLATED) as target:
            for info in source.infolist():
                data = source.read(info.filename)
                if info.filename.startswith("word/") and info.filename.endswith(".xml"):
                    text = data.decode("utf-8", errors="ignore")
                    for key, value in replacements.items():
                        text = text.replace("{{" + key + "}}", xml_escape(value))
                    text = re.sub(r"<w:highlight\b[^>]*/>", "", text)
                    text = re.sub(r"<w:updateFields\b[^>]*/>", "", text)
                    data = text.encode("utf-8")
                target.writestr(info, data)
        shutil.move(str(temp_path), str(path))
    finally:
        if temp_path.exists():
            temp_path.unlink()


def get_field_code_counts(path: Path) -> Dict[str, int]:
    counts = {"fldChar": 0, "TOC": 0, "PAGEREF": 0, "PAGE": 0}
    with ZipFile(path, "r") as zf:
        for name in zf.namelist():
            if not (name.startswith("word/") and name.endswith(".xml")):
                continue
            text = zf.read(name).decode("utf-8", errors="ignore")
            counts["fldChar"] += len(re.findall(r"<w:fldChar\b", text))
            for instruction in re.findall(
                r"<w:instrText\b[^>]*>(.*?)</w:instrText>",
                text,
                flags=re.DOTALL,
            ):
                normalized = re.sub(r"<[^>]+>", "", instruction).strip().upper()
                for field_name in ("TOC", "PAGEREF", "PAGE"):
                    if re.match(rf"^{field_name}\b", normalized):
                        counts[field_name] += 1
    return counts


def update_word_fields(path: Path) -> Optional[str]:
    if not sys.platform.startswith("win"):
        return "当前系统不是 Windows，已保留目录域，但未自动刷新目录页码"
    try:
        import pythoncom
        import win32com.client
    except Exception as exc:
        return f"未安装或无法加载 Word 自动化组件，已保留目录域但未自动刷新: {exc}"

    word = None
    doc = None
    pythoncom.CoInitialize()
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        try:
            word.AutomationSecurity = 3
        except Exception:
            pass
        doc = word.Documents.Open(
            str(path),
            ConfirmConversions=False,
            ReadOnly=False,
            AddToRecentFiles=False,
        )
        for story_type in range(1, 18):
            try:
                story_range = doc.StoryRanges(story_type)
            except Exception:
                continue
            while story_range is not None:
                try:
                    for field_index in range(
                        story_range.Fields.Count,
                        0,
                        -1,
                    ):
                        field = story_range.Fields.Item(field_index)
                        if field.Type != 13:  # wdFieldTOC
                            field.Update()
                except Exception:
                    pass
                try:
                    story_range = story_range.NextStoryRange
                except Exception:
                    story_range = None
        doc.Repaginate()
        for index in range(1, doc.TablesOfContents.Count + 1):
            doc.TablesOfContents.Item(index).UpdatePageNumbers()
        doc.Repaginate()
        doc.Save()
        return None
    except Exception as exc:
        return f"Word 后台更新目录/页码失败，域代码仍已保留: {exc}"
    finally:
        if doc is not None:
            try:
                doc.Close(SaveChanges=False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()


def scan_docx(
    path: Path,
    expected_field_counts: Optional[Dict[str, int]] = None,
) -> List[str]:
    issues: List[str] = []
    try:
        Document(path)
    except Exception as exc:
        issues.append(f"文件无法用 python-docx 打开: {exc}")
    with ZipFile(path, "r") as zf:
        for name in zf.namelist():
            if not (name.startswith("word/") and name.endswith(".xml")):
                continue
            text = zf.read(name).decode("utf-8", errors="ignore")
            if "{{" in text or "}}" in text:
                issues.append(f"{name} 中仍有占位符")
            if "<w:highlight" in text:
                issues.append(f"{name} 中仍有高亮")
            if "<w:updateFields" in text:
                issues.append(f"{name} 中仍有自动更新域设置，Word 打开时可能弹出提示")
    if expected_field_counts:
        actual_counts = get_field_code_counts(path)
        for field_name, expected in expected_field_counts.items():
            actual = actual_counts.get(field_name, 0)
            if actual < expected:
                issues.append(
                    f"Word 域代码数量减少: {field_name} 模板={expected}, 输出={actual}"
                )
    return issues


def render_report(
    template_file: Path,
    output_file: Path,
    report: PreparedReport,
    expected_field_counts: Dict[str, int],
    update_fields: bool,
) -> List[str]:
    doc = Document(template_file)
    delete_optional_blocks(doc, report)
    replace_placeholders(doc, report)
    add_guarantor_extra_lines(doc, report)
    clear_highlights(doc)
    output_file.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_file))
    postprocess_docx_xml(output_file, report.replacements)
    issues: List[str] = []
    if update_fields:
        field_update_warning = update_word_fields(output_file)
        if field_update_warning:
            issues.append(field_update_warning)
        postprocess_docx_xml(output_file, report.replacements)
    issues.extend(scan_docx(output_file, expected_field_counts))
    return issues


def find_single_file(directory: Path, pattern: str, description: str) -> Path:
    files = sorted(directory.glob(pattern))
    if not files:
        raise FileNotFoundError(f"未找到{description}: {directory / pattern}")
    if len(files) > 1:
        raise ValueError(f"找到多个{description}，请通过参数指定: {', '.join(str(p) for p in files)}")
    return files[0]


def get_runtime_project_root() -> Path:
    if not bool(getattr(sys, "frozen", False)):
        return Path(__file__).resolve().parent.parent
    executable_dir = Path(sys.executable).resolve().parent
    if (
        executable_dir.name.lower() == "dist"
        and (executable_dir.parent / "bin").is_dir()
    ):
        return executable_dir.parent
    return executable_dir


def resolve_defaults(project_root: Path) -> Dict[str, Path]:
    template_dir = project_root / "bin" / "template"
    resource_dir = Path(getattr(sys, "_MEIPASS", "")) / "resources"
    if bool(getattr(sys, "frozen", False)) and resource_dir.is_dir():
        template_file = resource_dir / "report_template.docx"
        prompt_file = resource_dir / "report_prompt.md"
        rules_file = resource_dir / "report_rules.json"
    else:
        template_file = find_single_file(
            template_dir,
            "*基底模板.docx",
            "Word模板",
        )
        prompt_file = find_single_file(
            template_dir,
            "*Prompt.md",
            "Prompt文件",
        )
        rules_file = find_single_file(
            template_dir,
            "*规则.json",
            "规则JSON",
        )
    return {
        "rules_file": rules_file,
        "template_file": template_file,
        "prompt_file": prompt_file,
        "records_file": project_root / "bin" / "json" / "图片提取数据.json",
    }


def load_rules(rules_file: Path) -> Dict[str, Any]:
    return json.loads(rules_file.read_text(encoding="utf-8-sig"))


def merge_saved_guarantors(
    rules: Dict[str, Any],
    guarantors_file: Path,
) -> Dict[str, Any]:
    merged = deepcopy(rules)
    if not guarantors_file.is_file():
        return merged
    payload = json.loads(guarantors_file.read_text(encoding="utf-8-sig"))
    guarantors = payload.get("guarantors", {})
    if not isinstance(guarantors, dict):
        raise ValueError(f"保证人资料的 guarantors 必须是对象: {guarantors_file}")
    merged["guarantors"] = guarantors
    return merged


def load_project_values(rules: Dict[str, Any]) -> Dict[str, Any]:
    project = dict(DEFAULT_PROJECT_VALUES)
    project.update(rules.get("project_defaults", {}))
    return project


def load_gemini_api_key(api_key_file: Path) -> str:
    environment_key = os.environ.get("GEMINI_API_KEY", "").strip()
    if environment_key:
        return environment_key
    if api_key_file.is_file():
        file_key = api_key_file.read_text(encoding="utf-8-sig").strip()
        if file_key:
            return file_key
    raise ValueError(
        "未找到 Gemini API Key。请在面板中填写，或设置 GEMINI_API_KEY，"
        f"或写入本机文件: {api_key_file}"
    )


def load_records(records_file: Path) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
    payload = json.loads(records_file.read_text(encoding="utf-8-sig"))
    if not isinstance(payload, dict):
        raise ValueError(f"图片提取 JSON 顶层必须是对象: {records_file}")
    raw_records = payload.get("records", [])
    extraction_errors = payload.get("errors", [])
    if not isinstance(raw_records, list):
        raise ValueError(f"图片提取 JSON 的 records 必须是数组: {records_file}")
    records: List[Dict[str, Any]] = []
    errors: List[Dict[str, Any]] = []
    for index, item in enumerate(raw_records, start=1):
        if not isinstance(item, dict):
            errors.append(
                {
                    "row_number": index,
                    "output_status": "error",
                    "error": "图片提取记录不是 JSON 对象",
                }
            )
            continue
        data = item.get("data")
        missing_fields = item.get("missing_fields") or []
        if not isinstance(data, dict):
            errors.append(
                {
                    "row_number": index,
                    "folder": clean_text(item.get("folder")),
                    "output_status": "error",
                    "error": "图片提取记录缺少 data 对象",
                }
            )
            continue
        if missing_fields:
            errors.append(
                {
                    "row_number": index,
                    "folder": clean_text(item.get("folder")),
                    "sequence": clean_text(data.get("序号")),
                    "name": clean_text(data.get("姓名")),
                    "output_status": "error",
                    "error": "图片数据缺少生成字段: "
                    + ", ".join(clean_text(field) for field in missing_fields),
                }
            )
            continue
        record = dict(data)
        record["_source_folder"] = clean_text(item.get("folder"))
        record["_source_images"] = item.get("images") or []
        record["_extraction_warnings"] = item.get("warnings") or []
        record["_gemini"] = item.get("gemini") or {}
        record["_excel"] = item.get("excel") or {}
        records.append(record)
    if isinstance(extraction_errors, list):
        for item in extraction_errors:
            if isinstance(item, dict):
                errors.append(
                    {
                        "folder": clean_text(item.get("folder")),
                        "output_status": "error",
                        "error": clean_text(item.get("error"))
                        or "图片提取失败",
                    }
                )
    return records, errors


def validate_document_type(rules: Dict[str, Any], document_type: str) -> None:
    config = rules.get("document_types", {}).get(document_type)
    if not config:
        raise ValueError(f"规则中未定义文档类型 {document_type}")
    if not config.get("enabled"):
        name = clean_text(config.get("name")) or "待配置"
        raise ValueError(f"文档类型 {document_type}（{name}）尚未配置模板和生成规则")


def build_parser(project_root: Path) -> argparse.ArgumentParser:
    defaults = resolve_defaults(project_root)
    parser = argparse.ArgumentParser(
        description="Generate one value-analysis Word report per image-extracted JSON record."
    )
    parser.add_argument("--template-file", default=str(defaults["template_file"]))
    parser.add_argument("--prompt-file", default=str(defaults["prompt_file"]))
    parser.add_argument("--rules-file", default=str(defaults["rules_file"]))
    parser.add_argument(
        "--guarantors-file",
        default=str(project_root / "bin" / "json" / "保证人资料.json"),
    )
    parser.add_argument("--records-file", default=str(defaults["records_file"]))
    parser.add_argument("--word-dir", default=str(project_root / "word"))
    parser.add_argument("--document-type", choices=["1", "2"], default="1")
    parser.add_argument("--model", default=None)
    parser.add_argument(
        "--api-key-file",
        default=str(project_root / "gemini_api.txt"),
        help="Local ignored file containing the Gemini API key.",
    )
    parser.add_argument(
        "--skip-gemini",
        action="store_true",
        help="Compatibility flag: do not run the optional per-row Gemini audit.",
    )
    parser.add_argument(
        "--use-gemini-audit",
        action="store_true",
        help="Optional: run the legacy per-row Gemini audit after GUI guarantor extraction.",
    )
    parser.add_argument(
        "--no-word-field-update",
        action="store_true",
        help="Preserve Word fields but skip Microsoft Word COM field refresh.",
    )
    return parser


def main(argv: Optional[Sequence[str]] = None) -> int:
    project_root = get_runtime_project_root()
    args = build_parser(project_root).parse_args(argv)

    template_file = Path(args.template_file).expanduser().resolve()
    prompt_file = Path(args.prompt_file).expanduser().resolve()
    rules_file = Path(args.rules_file).expanduser().resolve()
    guarantors_file = Path(args.guarantors_file).expanduser().resolve()
    records_file = Path(args.records_file).expanduser().resolve()
    word_dir = Path(args.word_dir).expanduser().resolve()
    api_key_file = Path(args.api_key_file).expanduser().resolve()

    for label, path in [
        ("Word模板", template_file),
        ("Prompt文件", prompt_file),
        ("规则JSON", rules_file),
        ("图片提取JSON", records_file),
    ]:
        if not path.is_file():
            raise FileNotFoundError(f"{label}不存在: {path}")

    # The prompt is part of the configured generation package; reading it here
    # catches missing/encoding issues before any report is written.
    prompt_text = prompt_file.read_text(encoding="utf-8-sig").strip()
    if not prompt_text:
        raise ValueError(f"Prompt文件为空: {prompt_file}")

    rules = merge_saved_guarantors(
        load_rules(rules_file),
        guarantors_file,
    )
    validate_document_type(rules, args.document_type)
    project = load_project_values(rules)
    ai_config = rules.get("ai", {})
    model = clean_text(args.model) or clean_text(ai_config.get("model")) or DEFAULT_GEMINI_MODEL
    thinking_level = clean_text(ai_config.get("thinking_level")) or "medium"
    use_google_search = bool(ai_config.get("use_google_search_for_guarantor", False))
    store_interactions = bool(ai_config.get("store_interactions", False))
    saved_guarantors = rules.get("guarantors", {})
    gemini_enabled = bool(args.use_gemini_audit and not args.skip_gemini)
    gemini_client = None
    if gemini_enabled:
        gemini_client = create_gemini_client(load_gemini_api_key(api_key_file))

    records, extraction_errors = load_records(records_file)
    template_field_counts = get_field_code_counts(template_file)
    expected_field_counts = {
        field_name: template_field_counts[field_name]
        for field_name in ("TOC", "PAGEREF", "PAGE")
    }

    manifest: Dict[str, Any] = {
        "document_type": args.document_type,
        "template": str(template_file),
        "prompt": str(prompt_file),
        "rules": str(rules_file),
        "guarantors_file": str(guarantors_file),
        "records_file": str(records_file),
        "source_type": "image_folders_with_excel",
        "ai": {
            "provider": "Google Gemini",
            "model": model,
            "api": "Interactions API",
            "enabled": gemini_enabled,
            "store_interactions": store_interactions,
        },
        "template_field_counts": template_field_counts,
        "reports": [],
        "errors": list(extraction_errors),
    }

    total_records = 0
    generated = 0
    print(f"Using template: {template_file}")
    print(f"Using prompt: {prompt_file}")
    print(f"Using image records: {records_file}")
    print(f"Document type: {args.document_type}")
    print(
        f"AI audit: {'disabled' if not gemini_enabled else f'Google Gemini / {model} / Interactions API'}"
    )

    for row_number, values in enumerate(records, start=1):
        total_records += 1
        sequence = clean_text(values.get("序号"))
        name = clean_text(values.get("姓名"))
        source_folder = clean_text(values.get("_source_folder"))
        print(
            f"Processing record {row_number} | 序号 {sequence} | "
            f"{name} | {source_folder}"
        )
        try:
            guarantor_name = clean_text(values.get("保证人"))
            saved_guarantor_details = (
                dict(saved_guarantors.get(guarantor_name, {}))
                if guarantor_name
                else {}
            )
            saved_guarantor_details.pop("保证人", None)
            excel_guarantor_details = get_excel_guarantor_details(values)
            guarantor_details = dict(saved_guarantor_details)
            guarantor_details.update(excel_guarantor_details)
            validate_guarantor_details(guarantor_name, guarantor_details)
            report = build_prepared_report(
                row_number,
                values,
                project,
                guarantor_details,
            )
            report.warnings.extend(
                f"图片提取: {clean_text(warning)}"
                for warning in values.get("_extraction_warnings", [])
                if clean_text(warning)
            )
            audit: Optional[GeminiAudit] = None
            if gemini_client is not None:
                print(f"  Calling Gemini: {model}")
                audit = call_gemini(
                    client=gemini_client,
                    model=model,
                    thinking_level=thinking_level,
                    prompt_text=prompt_text,
                    values=values,
                    report=report,
                    excel_details=guarantor_details,
                    use_google_search=use_google_search,
                    store_interactions=store_interactions,
                )
                merge_gemini_guarantor_details(report, excel_guarantor_details, audit)
            output_file = word_dir / report.filename
            validation_issues = render_report(
                template_file,
                output_file,
                report,
                expected_field_counts=expected_field_counts,
                update_fields=not args.no_word_field_update,
            )
            all_warnings = report.warnings + validation_issues
            report_manifest = {
                "row_number": report.row_number,
                "sequence": report.sequence,
                "branch": report.branch,
                "name": report.name,
                "source_folder": source_folder,
                "source_images": values.get("_source_images", []),
                "image_extraction": values.get("_gemini", {}),
                "excel_source": values.get("_excel", {}),
                "filename": report.filename,
                "report_number": report.report_number,
                "debtor_count": report.debtor_count,
                "guarantor_included": report.guarantor_included,
                "placement_cost_included": report.placement_cost_included,
                "minimum_item": report.minimum_item,
                "output_status": "success" if not all_warnings else "warning",
                "warnings": all_warnings,
                "word_field_counts": get_field_code_counts(output_file),
            }
            if audit is not None:
                report_manifest["gemini"] = {
                    "interaction_id": audit.interaction_id,
                    "model": audit.model,
                    "usage": audit.usage,
                    "guarantor_sources": audit.sources,
                }
            manifest["reports"].append(report_manifest)
            generated += 1
            print(f"  Saved: {output_file}")
            for warning in all_warnings:
                print(f"  WARNING: {warning}")
        except Exception as exc:
            error = {
                "row_number": row_number,
                "sequence": sequence,
                "name": name,
                "output_status": "error",
                "error": str(exc),
            }
            manifest["errors"].append(error)
            print(f"  ERROR: {exc}")

    word_dir.mkdir(parents=True, exist_ok=True)
    manifest_file = word_dir / "generation_manifest.json"
    manifest_file.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

    print(f"\nRows selected: {total_records}")
    print(f"Reports generated: {generated}")
    print(f"Errors: {len(manifest['errors'])}")
    print(f"Manifest: {manifest_file}")
    return 0 if not manifest["errors"] else 1


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        raise SystemExit(1)
