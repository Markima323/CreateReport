#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import os
import re
import sys
from collections import Counter
from contextlib import contextmanager
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, Dict, Iterable, Iterator, List, Optional, Sequence, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

from process_excel_to_word import (
    create_gemini_client,
    get_interaction_output_text,
    load_gemini_api_key,
    update_word_fields,
)


for stream in (sys.stdout, sys.stderr):
    if hasattr(stream, "reconfigure"):
        stream.reconfigure(encoding="utf-8", errors="replace")


SUPPORTED_EXCEL_SUFFIXES = {".xls", ".xlsx", ".xlsm"}
SUMMARY_SHEET = "汇总-资产基础法"
CLASSIFICATION_SHEET = "2分类汇总"
DEFAULT_MODEL = "gemini-3.5-flash"

DEFAULT_WORKPAPER_TABLE_OPTIONS: Dict[str, Any] = {
    "enabled": True,
    "insert_in_explanation_technical": True,
    "send_to_ai": False,
    "max_scan_rows": 500,
    "max_scan_columns": 40,
    "max_rows_per_table": 80,
    "max_columns_per_table": 40,
    "max_tables_per_workpaper": 12,
    "max_tables_per_subject": 10,
}

WORKPAPER_TABLE_KEYWORDS = (
    "案例",
    "实例",
    "可比",
    "比较因素",
    "交易价格",
    "交易情况",
    "市场状况",
    "修正",
    "系数",
    "参数",
    "取费",
    "指数",
    "重置",
    "成新",
    "耐用年限",
    "尚可使用年限",
    "经济寿命",
    "折旧",
    "摊销",
    "租金",
    "收益",
    "资本化率",
    "折现率",
    "净现金流",
    "市盈率",
    "市净率",
    "P/B",
    "P/E",
    "比率",
    "折扣率",
    "溢价率",
    "权重",
    "财务指标",
    "非流动性折扣",
    "流动性折扣",
    "控制权溢价",
    "少数股权",
    "基本重置价",
    "层高调整",
    "土地使用年期",
)

WORKPAPER_HEADER_KEYWORDS = (
    "序号",
    "项目",
    "指标",
    "因素",
    "案例",
    "实例",
    "评估对象",
    "公司名称",
    "名称",
    "交易价格",
    "比 较 因 素",
)

WORKPAPER_EXCLUDE_TABLE_KEYWORDS = (
    "资产负债表--资产",
    "资产负债表--负债",
)

WORKPAPER_SUBJECT_HINTS: Sequence[Tuple[str, Sequence[str]]] = (
    (
        "长期股权投资",
        (
            "长期股权",
            "股权",
            "企业价值",
            "交易案例",
            "对比公司",
            "上市公司",
            "财务指标",
            "中小企业发展指数",
            "非流动性折扣",
            "控制权溢价",
            "少数股权",
        ),
    ),
    ("房屋建筑物", ("房屋", "建筑物", "厂房", "办公楼", "不动产", "租金")),
    ("构筑物及其他辅助设施", ("构筑物", "道路", "围墙", "冷却塔")),
    ("土地使用权", ("土地", "地价", "宗地", "土地使用年期", "工业用地")),
    ("机器设备", ("机器设备", "生产设备", "设备评估", "通用设备", "专用设备")),
    ("车辆", ("车辆", "运输设备")),
    ("电子设备", ("电子设备", "办公设备", "电脑")),
    ("存货", ("存货", "原材料", "产成品", "库存商品", "发出商品")),
    ("在建工程", ("在建工程",)),
    ("技术类无形资产", ("专利", "专有技术", "软件著作权")),
    ("其他无形资产", ("商标", "软件", "无形资产")),
    ("投资性房地产", ("投资性房地产", "商铺", "租金")),
)

RESPONSE_SCHEMA: Dict[str, Any] = {
    "type": "object",
    "properties": {
        "report": {
            "type": "object",
            "properties": {
                "scope_intro": {"type": "array", "items": {"type": "string"}},
                "asset_descriptions": {
                    "type": "array",
                    "items": {"type": "string"},
                },
                "method_sections": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "subject": {"type": "string"},
                            "paragraphs": {
                                "type": "array",
                                "items": {"type": "string"},
                            },
                        },
                        "required": ["subject", "paragraphs"],
                    },
                },
                "conclusion_intro": {
                    "type": "array",
                    "items": {"type": "string"},
                },
                "final_conclusion": {
                    "type": "array",
                    "items": {"type": "string"},
                },
            },
            "required": [
                "scope_intro",
                "asset_descriptions",
                "method_sections",
                "conclusion_intro",
                "final_conclusion",
            ],
        },
        "explanation": {
            "type": "object",
            "properties": {
                "scope_intro": {"type": "array", "items": {"type": "string"}},
                "physical_assets": {
                    "type": "array",
                    "items": {"type": "string"},
                },
                "technical_sections": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "subject": {"type": "string"},
                            "paragraphs": {
                                "type": "array",
                                "items": {"type": "string"},
                            },
                        },
                        "required": ["subject", "paragraphs"],
                    },
                },
            },
            "required": [
                "scope_intro",
                "physical_assets",
                "technical_sections",
            ],
        },
        "warnings": {"type": "array", "items": {"type": "string"}},
    },
    "required": ["report", "explanation", "warnings"],
}


def clean_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return re.sub(r"\s+", " ", str(value)).strip()


def normalize_name(value: Any) -> str:
    return re.sub(r"[\s　:：()（）\-—_]+", "", clean_text(value)).lower()


def normalize_sheet_name(value: Any) -> str:
    return re.sub(r"\s+", "", clean_text(value))


def normalize_matrix(value: Any) -> List[List[Any]]:
    if value is None:
        return []
    if not isinstance(value, tuple):
        return [[value]]
    if value and not isinstance(value[0], tuple):
        return [list(value)]
    return [list(row) for row in value]


def numeric_value(value: Any) -> Optional[float]:
    if value in (None, ""):
        return None
    if isinstance(value, (int, float)):
        return float(value)
    text = clean_text(value).replace(",", "").replace("%", "")
    if not text or text in {"-", "—", "/"}:
        return None
    try:
        number = float(text)
    except ValueError:
        return None
    if "%" in clean_text(value):
        number /= 100
    return number


def format_number(value: Any, decimals: int = 2) -> str:
    number = numeric_value(value)
    if number is None:
        return "-"
    return f"{number:,.{decimals}f}"


def format_rate(value: Any) -> str:
    number = numeric_value(value)
    if number is None:
        return "-"
    return f"{number * 100:.2f}%"


def format_excel_date_serial(value: Any) -> str:
    number = numeric_value(value)
    if number is None or not (20000 <= number <= 60000):
        return ""
    date_value = datetime(1899, 12, 30) + timedelta(days=int(number))
    return f"{date_value.year}年{date_value.month}月{date_value.day}日"


def format_table_cell(value: Any, header: str = "") -> str:
    if value in (None, ""):
        return ""
    header_name = clean_text(header)
    number = numeric_value(value)
    if number is None:
        return clean_text(value)
    if any(marker in header_name for marker in ("日期", "年月", "启用", "到期")):
        date_text = format_excel_date_serial(number)
        if date_text:
            return date_text
    if "率" in header_name and "利率" not in header_name:
        return format_rate(number)
    if any(marker in header_name for marker in ("账号", "号码", "编号", "证号")):
        return str(int(number)) if number.is_integer() else clean_text(value)
    if any(
        marker in header_name
        for marker in (
            "金额",
            "价值",
            "原值",
            "净值",
            "单价",
            "注册资本",
            "实收资本",
            "出资额",
            "面积",
            "数量",
        )
    ):
        return format_number(number)
    if number.is_integer():
        return str(int(number))
    return clean_text(value)


def sanitize_filename_part(value: str) -> str:
    cleaned = re.sub(r'[\\/:*?"<>|]', "_", clean_text(value))
    return cleaned.rstrip(". ") or "未命名"


def discover_excel_files(input_dir: Path) -> List[Path]:
    files = sorted(
        path
        for path in input_dir.rglob("*")
        if (
            path.is_file()
            and not path.name.startswith("~$")
            and path.suffix.lower() in SUPPORTED_EXCEL_SUFFIXES
        )
    )
    if not files:
        raise FileNotFoundError(f"项目目录中未找到 Excel 文件: {input_dir}")
    return files


def discover_project_folders(input_root: Path) -> List[Path]:
    if not input_root.is_dir():
        raise FileNotFoundError(f"输入根目录不存在: {input_root}")
    projects = []
    for folder in sorted(
        (path for path in input_root.iterdir() if path.is_dir()),
        key=lambda path: path.name,
    ):
        has_detail_workbook = any(
            path.is_file()
            and not path.name.startswith("~$")
            and path.suffix.lower() in SUPPORTED_EXCEL_SUFFIXES
            and "评估明细表" in path.name
            for path in folder.rglob("*")
        )
        if has_detail_workbook:
            projects.append(folder)
    if projects:
        return projects

    root_is_project = any(
        path.is_file()
        and not path.name.startswith("~$")
        and path.suffix.lower() in SUPPORTED_EXCEL_SUFFIXES
        and "评估明细表" in path.name
        for path in input_root.rglob("*")
    )
    if root_is_project:
        return [input_root]
    raise ValueError(
        "输入根目录中没有类型 2 项目子文件夹；"
        "每个项目文件夹至少应包含一份名称带“评估明细表”的 Excel"
    )


@contextmanager
def excel_application() -> Iterator[Any]:
    if not sys.platform.startswith("win"):
        raise RuntimeError("类型 2 的 .xls 读取目前需要 Windows 和 Microsoft Excel")
    import pythoncom
    import win32com.client

    pythoncom.CoInitialize()
    excel = None
    try:
        excel = win32com.client.DispatchEx("Excel.Application")
        excel.Visible = False
        excel.DisplayAlerts = False
        excel.EnableEvents = False
        excel.AskToUpdateLinks = False
        try:
            excel.AutomationSecurity = 3
        except Exception:
            pass
        yield excel
    finally:
        if excel is not None:
            excel.Quit()
        pythoncom.CoUninitialize()


@contextmanager
def open_workbook(excel: Any, path: Path) -> Iterator[Any]:
    workbook = excel.Workbooks.Open(
        str(path),
        UpdateLinks=0,
        ReadOnly=True,
        IgnoreReadOnlyRecommended=True,
        AddToMru=False,
        Notify=False,
    )
    try:
        yield workbook
    finally:
        workbook.Close(False)


def visible_worksheets(workbook: Any) -> List[Any]:
    return [
        workbook.Worksheets.Item(index)
        for index in range(1, workbook.Worksheets.Count + 1)
        if int(workbook.Worksheets.Item(index).Visible) == -1
    ]


def find_worksheet(workbook: Any, expected_name: str) -> Optional[Any]:
    expected = normalize_sheet_name(expected_name)
    for index in range(1, workbook.Worksheets.Count + 1):
        sheet = workbook.Worksheets.Item(index)
        if normalize_sheet_name(sheet.Name) == expected:
            return sheet
    return None


def read_range(
    sheet: Any,
    *,
    max_rows: Optional[int] = None,
    max_columns: Optional[int] = None,
) -> List[List[Any]]:
    used = sheet.UsedRange
    first_row = int(used.Row)
    first_column = int(used.Column)
    row_count = int(used.Rows.Count)
    column_count = int(used.Columns.Count)
    if max_rows is not None:
        row_count = min(row_count, max_rows)
    if max_columns is not None:
        column_count = min(column_count, max_columns)
    if row_count < 1 or column_count < 1:
        return []
    value = sheet.Range(
        sheet.Cells(first_row, first_column),
        sheet.Cells(first_row + row_count - 1, first_column + column_count - 1),
    ).Value2
    return normalize_matrix(value)


def find_prefixed_value(matrix: Sequence[Sequence[Any]], label: str) -> str:
    normalized_label = normalize_name(label)
    for row in matrix[:12]:
        for cell in row:
            text = clean_text(cell)
            compact = normalize_name(text)
            if normalized_label not in compact:
                continue
            for separator in ("：", ":"):
                if separator in text:
                    return text.split(separator, 1)[1].strip()
    return ""


def find_header_row(
    matrix: Sequence[Sequence[Any]],
    required_headers: Sequence[str],
) -> int:
    required = [normalize_name(header) for header in required_headers]
    for index, row in enumerate(matrix[:20]):
        cells = [normalize_name(cell) for cell in row]
        if all(any(header in cell for cell in cells) for header in required):
            return index
    return -1


def header_column(row: Sequence[Any], expected: str) -> int:
    target = normalize_name(expected)
    for index, cell in enumerate(row):
        if target in normalize_name(cell):
            return index
    return -1


def row_value(row: Sequence[Any], index: int) -> Any:
    return row[index] if 0 <= index < len(row) else None


def parse_summary_matrix(matrix: Sequence[Sequence[Any]]) -> Dict[str, Any]:
    entity = find_prefixed_value(matrix, "被评估单位")
    benchmark_date = find_prefixed_value(matrix, "评估基准日")
    header_index = find_header_row(matrix, ("账面价值", "评估价值", "增减值"))
    if header_index < 0:
        raise ValueError(f"{SUMMARY_SHEET} 未找到汇总表头")
    header = matrix[header_index]
    book_col = header_column(header, "账面价值")
    assessed_col = header_column(header, "评估价值")
    change_col = header_column(header, "增减值")
    rate_col = header_column(header, "增值率")
    rows = []
    for source_row in matrix[header_index + 1 :]:
        sequence = clean_text(row_value(source_row, 0))
        project = clean_text(row_value(source_row, 1))
        if not project and source_row:
            candidate = clean_text(source_row[0])
            if candidate and not re.fullmatch(r"\d+", candidate):
                project = candidate
        if not project or project in {"A", "B", "C=B-A", "D=C/A×100%"}:
            continue
        book = numeric_value(row_value(source_row, book_col))
        assessed = numeric_value(row_value(source_row, assessed_col))
        change = numeric_value(row_value(source_row, change_col))
        rate = numeric_value(row_value(source_row, rate_col))
        if all(value is None for value in (book, assessed, change, rate)):
            continue
        rows.append(
            {
                "序号": sequence,
                "项目": project,
                "账面价值": book,
                "评估价值": assessed,
                "增减值": change,
                "增值率": rate,
            }
        )
    if not entity:
        raise ValueError(f"{SUMMARY_SHEET} 未找到被评估单位")
    return {
        "sheet": SUMMARY_SHEET,
        "unit": "人民币万元",
        "entity": entity,
        "benchmark_date": benchmark_date,
        "rows": rows,
    }


def parse_classification_matrix(
    matrix: Sequence[Sequence[Any]],
) -> Dict[str, Any]:
    header_index = find_header_row(matrix, ("科目名称", "账面价值", "评估价值"))
    if header_index < 0:
        raise ValueError(f"{CLASSIFICATION_SHEET} 未找到分类汇总表头")
    header = matrix[header_index]
    subject_col = header_column(header, "科目名称")
    book_col = header_column(header, "账面价值")
    assessed_col = header_column(header, "评估价值")
    change_col = header_column(header, "增减值")
    rate_col = header_column(header, "增值率")
    rows = []
    category = ""
    for source_row in matrix[header_index + 1 :]:
        sequence = clean_text(row_value(source_row, 0))
        category_marker = clean_text(row_value(source_row, 1))
        subject = clean_text(row_value(source_row, subject_col))
        if not subject:
            continue
        if category_marker:
            category = subject.replace("合计", "").strip()
        rows.append(
            {
                "序号": sequence,
                "类别": category,
                "科目名称": subject.replace("合计", "").strip(),
                "账面价值": numeric_value(row_value(source_row, book_col)),
                "评估价值": numeric_value(row_value(source_row, assessed_col)),
                "增减值": numeric_value(row_value(source_row, change_col)),
                "增值率": numeric_value(row_value(source_row, rate_col)),
            }
        )
    return {
        "sheet": CLASSIFICATION_SHEET,
        "unit": "人民币元",
        "rows": rows,
    }


def read_detail_candidate(excel: Any, path: Path) -> Dict[str, Any]:
    with open_workbook(excel, path) as workbook:
        summary_sheet = find_worksheet(workbook, SUMMARY_SHEET)
        classification_sheet = find_worksheet(workbook, CLASSIFICATION_SHEET)
        if summary_sheet is None or classification_sheet is None:
            raise ValueError("缺少汇总-资产基础法或2分类汇总")
        summary = parse_summary_matrix(
            read_range(summary_sheet, max_rows=120, max_columns=12)
        )
        classification = parse_classification_matrix(
            read_range(classification_sheet, max_rows=180, max_columns=12)
        )
        return {
            "file": str(path),
            "entity": summary["entity"],
            "benchmark_date": summary["benchmark_date"],
            "summary": summary,
            "classification": classification,
        }


def selection_score(input_dir: Path, candidate: Dict[str, Any]) -> Tuple[int, str]:
    path = Path(candidate["file"])
    entity = clean_text(candidate["entity"])
    root_name = normalize_name(input_dir.name)
    entity_name = normalize_name(entity)
    filename = normalize_name(path.stem)
    score = 0
    reasons = []
    if entity_name and (entity_name in root_name or root_name in entity_name):
        score += 100
        reasons.append("被评估单位与项目目录名称一致")
    if re.match(r"^3(?:-|_)", path.name):
        score += 40
        reasons.append("文件名以3-开头")
    elif re.match(r"^3\.1(?:-|_)", path.name):
        score += 35
        reasons.append("文件名以3.1-开头")
    if entity_name and any(
        token and token in filename
        for token in re.split(r"公司|股份|有限|科技|集团", entity_name)
    ):
        score += 15
        reasons.append("文件名包含被评估单位简称")
    return score, "；".join(reasons) or "按文件名顺序候选"


def choose_primary_workbook(
    input_dir: Path,
    candidates: Sequence[Dict[str, Any]],
) -> Dict[str, Any]:
    if not candidates:
        raise ValueError("未找到同时包含汇总-资产基础法和2分类汇总的评估明细表")
    ranked = []
    for candidate in candidates:
        score, reason = selection_score(input_dir, candidate)
        ranked.append((score, reason, candidate))
    ranked.sort(key=lambda item: (-item[0], item[2]["file"]))
    if len(ranked) > 1 and ranked[0][0] == ranked[1][0]:
        first = Path(ranked[0][2]["file"]).name
        second = Path(ranked[1][2]["file"]).name
        raise ValueError(
            "无法唯一确定母公司评估明细表，请调整文件名或目录名: "
            f"{first}, {second}"
        )
    primary = dict(ranked[0][2])
    primary["selection_reason"] = ranked[0][1]
    return primary


def company_aliases(candidate: Dict[str, Any]) -> List[str]:
    aliases = []
    entity = clean_text(candidate.get("entity"))
    entity_base = re.sub(
        r"(股份有限公司|有限责任公司|有限公司)$",
        "",
        entity,
    )
    for suffix in ("科技", "智慧科技", "无线科技", "投资", "集团"):
        if entity_base.endswith(suffix):
            entity_base = entity_base[: -len(suffix)]
            break
    if len(entity_base) >= 2:
        aliases.append(entity_base)

    stem = Path(clean_text(candidate.get("file"))).stem
    if "评估明细表-" in stem:
        stem = stem.split("评估明细表-", 1)[1]
    stem = re.sub(r"\d{4}(?:[.\-]\d+)*.*$", "", stem)
    stem = re.sub(r"(反馈修改稿|报告日.*)$", "", stem)
    stem = stem.strip("-_ ")
    if len(stem) >= 2 and stem not in aliases:
        aliases.append(stem)
    return aliases


def subject_from_sheet_name(name: str, aliases: Dict[str, str]) -> str:
    subject = re.sub(r"^\d+(?:\.\d+)*[.、\-]?", "", clean_text(name))
    subject = re.sub(r"(汇总表?|评估明细表)$", "", subject).strip()
    for alias, canonical in aliases.items():
        if normalize_name(alias) == normalize_name(subject):
            return canonical
    return subject


def combined_headers(
    matrix: Sequence[Sequence[Any]],
    header_index: int,
) -> List[str]:
    width = max((len(row) for row in matrix[: header_index + 2]), default=0)
    headers = []
    for column in range(width):
        parts = []
        for row_index in (header_index, header_index + 1):
            if row_index >= len(matrix):
                continue
            text = clean_text(row_value(matrix[row_index], column))
            if text and text not in parts:
                parts.append(text)
        headers.append("/".join(parts))
    return headers


def matching_column(headers: Sequence[str], candidates: Sequence[str]) -> int:
    for candidate in candidates:
        target = normalize_name(candidate)
        for index, header in enumerate(headers):
            if target in normalize_name(header):
                return index
    return -1


def is_detail_data_marker(value: Any) -> bool:
    text = clean_text(value)
    return bool(re.fullmatch(r"\d+(?:\.0+)?", text)) or text in {
        "合计",
        "小计",
        "总计",
    }


def first_detail_data_row(
    matrix: Sequence[Sequence[Any]],
    header_index: int,
    sequence_col: int,
) -> int:
    for index in range(header_index + 1, min(len(matrix), header_index + 8)):
        if is_detail_data_marker(row_value(matrix[index], sequence_col)):
            return index
    return min(header_index + 1, len(matrix))


def detail_table_headers(
    matrix: Sequence[Sequence[Any]],
    header_index: int,
    data_index: int,
) -> List[str]:
    width = max((len(row) for row in matrix[: max(data_index, header_index + 1)]), default=0)
    headers = []
    for column in range(width):
        parts = []
        for row_index in range(header_index, data_index):
            text = clean_text(row_value(matrix[row_index], column))
            if text and text not in parts:
                parts.append(text)
        headers.append("/".join(parts))
    return headers


def trim_table_columns(
    headers: Sequence[str],
    rows: Sequence[Sequence[Any]],
) -> Tuple[List[str], List[List[Any]], List[int]]:
    width = max([len(headers), *(len(row) for row in rows)], default=0)
    used_columns = []
    for column in range(width):
        header = clean_text(row_value(headers, column))
        has_data = any(clean_text(row_value(row, column)) for row in rows)
        if header or has_data:
            used_columns.append(column)
    if not used_columns:
        return [], [], []
    trimmed_headers = [
        clean_text(row_value(headers, column)) or f"列{index + 1}"
        for index, column in enumerate(used_columns)
    ]
    trimmed_rows = [
        [row_value(row, column) for column in used_columns]
        for row in rows
    ]
    return trimmed_headers, trimmed_rows, used_columns


def needs_display_text(header: str) -> bool:
    return any(
        marker in clean_text(header)
        for marker in (
            "账号",
            "号码",
            "编号",
            "证号",
            "权证",
        )
    )


def cell_display_text(
    sheet: Any,
    row_index: int,
    column_index: int,
) -> str:
    try:
        used = sheet.UsedRange
        first_row = int(used.Row)
        first_column = int(used.Column)
        return clean_text(
            sheet.Cells(
                first_row + row_index,
                first_column + column_index,
            ).Text
        )
    except Exception:
        return ""


def extract_detail_table(
    sheet: Any,
    matrix: Sequence[Sequence[Any]],
    header_index: int,
    subject: str,
    sheet_name: str,
) -> Dict[str, Any]:
    base_headers = combined_headers(matrix, header_index)
    sequence_col = matching_column(base_headers, ("序号",))
    if sequence_col < 0:
        return {
            "subject": subject,
            "sheet": sheet_name,
            "headers": [],
            "rows": [],
            "row_count": 0,
            "warnings": ["未找到序号列，未生成完整明细表"],
        }
    data_index = first_detail_data_row(matrix, header_index, sequence_col)
    headers = detail_table_headers(matrix, header_index, data_index)
    indexed_rows = [
        (source_index, row)
        for source_index, row in enumerate(matrix[data_index:], start=data_index)
        if any(clean_text(cell) for cell in row)
    ]
    headers, raw_rows, used_columns = trim_table_columns(
        headers,
        [row for _source_index, row in indexed_rows],
    )
    formatted_rows = []
    for (source_index, _source_row), row in zip(indexed_rows, raw_rows):
        formatted_row = []
        for index, value in enumerate(row):
            header = row_value(headers, index)
            if needs_display_text(header) and isinstance(value, (int, float)):
                display = cell_display_text(
                    sheet,
                    source_index,
                    row_value(used_columns, index) or 0,
                )
                if display and display != "###":
                    formatted_row.append(display)
                    continue
            formatted_row.append(format_table_cell(value, header))
        formatted_rows.append(formatted_row)
    return {
        "subject": subject,
        "sheet": sheet_name,
        "headers": headers,
        "rows": formatted_rows,
        "row_count": len(formatted_rows),
        "warnings": [],
    }


def summarize_subject_sheet(
    sheet: Any,
    aliases: Dict[str, str],
) -> Dict[str, Any]:
    matrix = read_range(sheet, max_rows=10000, max_columns=50)
    subject = subject_from_sheet_name(clean_text(sheet.Name), aliases)
    header_index = find_header_row(matrix, ("序号",))
    if header_index < 0:
        return {
            "subject": subject,
            "sheet": clean_text(sheet.Name),
            "item_count": None,
            "detail_headers": [],
            "detail_table": {
                "subject": subject,
                "sheet": clean_text(sheet.Name),
                "headers": [],
                "rows": [],
                "row_count": 0,
                "warnings": ["未找到序号表头"],
            },
            "content_summary": [],
            "location_summary": [],
            "representative_rows": [],
            "warnings": ["未找到序号表头"],
        }
    headers = combined_headers(matrix, header_index)
    detail_table = extract_detail_table(
        sheet,
        matrix,
        header_index,
        subject,
        clean_text(sheet.Name),
    )
    sequence_col = matching_column(headers, ("序号",))
    content_col = matching_column(
        headers,
        (
            "业务内容",
            "设备名称",
            "名称及规格型号",
            "名称",
            "证书名称及证书号",
            "欠款单位名称",
        ),
    )
    location_col = matching_column(
        headers,
        ("存放地点", "坐落", "位置", "使用地点"),
    )
    age_col = matching_column(headers, ("账龄", "库龄"))
    content_counts: Counter[str] = Counter()
    location_counts: Counter[str] = Counter()
    age_counts: Counter[str] = Counter()
    max_sequence = 0
    representative_rows = []
    for row in matrix[header_index + 1 :]:
        sequence_text = clean_text(row_value(row, sequence_col))
        match = re.fullmatch(r"(\d+)(?:\.0+)?", sequence_text)
        if not match:
            continue
        sequence = int(match.group(1))
        max_sequence = max(max_sequence, sequence)
        content = clean_text(row_value(row, content_col))
        location = clean_text(row_value(row, location_col))
        age = clean_text(row_value(row, age_col))
        if content:
            content_counts[content] += 1
        if location:
            location_counts[location] += 1
        if age:
            age_counts[age] += 1
        if len(representative_rows) < 8:
            representative_rows.append(
                {
                    headers[index] or f"列{index + 1}": clean_text(value)
                    for index, value in enumerate(row[: min(len(row), len(headers))])
                    if clean_text(value)
                }
            )
    return {
        "subject": subject,
        "sheet": clean_text(sheet.Name),
        "item_count": max_sequence or None,
        "detail_headers": [header for header in headers if header],
        "detail_table": detail_table,
        "content_summary": [
            {"value": value, "count": count}
            for value, count in content_counts.most_common(8)
        ],
        "location_summary": [
            {"value": value, "count": count}
            for value, count in location_counts.most_common(6)
        ],
        "main_age_bucket": age_counts.most_common(1)[0][0] if age_counts else "",
        "representative_rows": representative_rows,
        "warnings": [],
    }


def summarize_primary_workbook(
    excel: Any,
    primary: Dict[str, Any],
    rules: Dict[str, Any],
) -> Dict[str, Any]:
    path = Path(primary["file"])
    aliases = rules.get("subject_aliases", {})
    classification_by_subject = {
        normalize_name(row["科目名称"]): row
        for row in primary["classification"]["rows"]
    }
    with open_workbook(excel, path) as workbook:
        visible = visible_worksheets(workbook)
        hidden_names = [
            clean_text(workbook.Worksheets.Item(index).Name)
            for index in range(1, workbook.Worksheets.Count + 1)
            if int(workbook.Worksheets.Item(index).Visible) != -1
        ]
        subjects = []
        for sheet in visible:
            sheet_name = normalize_sheet_name(sheet.Name)
            if sheet_name in {
                normalize_sheet_name(SUMMARY_SHEET),
                normalize_sheet_name(CLASSIFICATION_SHEET),
            }:
                continue
            if "汇总" in sheet_name or sheet_name == "基本情况":
                continue
            summary = summarize_subject_sheet(sheet, aliases)
            classification = classification_by_subject.get(
                normalize_name(summary["subject"]),
                {},
            )
            summary.update(
                {
                    "category": classification.get("类别", ""),
                    "unit": "人民币元",
                    "book_value": classification.get("账面价值"),
                    "assessed_value": classification.get("评估价值"),
                    "increase_decrease": classification.get("增减值"),
                    "increase_rate": classification.get("增值率"),
                }
            )
            subjects.append(summary)
        return {
            "file": str(path),
            "entity": primary["entity"],
            "selection_reason": primary["selection_reason"],
            "visible_sheets": [clean_text(sheet.Name) for sheet in visible],
            "hidden_sheets_ignored": hidden_names,
            "subjects": subjects,
        }


def detect_method(texts: Iterable[str]) -> str:
    joined = "\n".join(texts)
    methods = [
        method
        for method in (
            "假设开发法",
            "上市公司比较法",
            "交易案例比较法",
            "市场法",
            "收益法",
            "重置成本法",
            "成本法",
            "资产基础法",
        )
        if method in joined
    ]
    return "、".join(methods)


def summarize_workpaper(excel: Any, path: Path) -> Dict[str, Any]:
    with open_workbook(excel, path) as workbook:
        sheets = visible_worksheets(workbook)
        key_text = []
        sheet_names = []
        for sheet in sheets:
            sheet_names.append(clean_text(sheet.Name))
            matrix = read_range(sheet, max_rows=15, max_columns=20)
            for row in matrix:
                for cell in row:
                    text = clean_text(cell)
                    if text and text not in key_text:
                        key_text.append(text)
                    if len(key_text) >= 80:
                        break
                if len(key_text) >= 80:
                    break
            if len(key_text) >= 80:
                break
        return {
            "file": str(path),
            "kind": detect_method([path.name, *sheet_names, *key_text]),
            "visible_sheets": sheet_names,
            "valuation_method": detect_method(key_text),
            "key_text": key_text[:40],
            "warnings": [],
        }


def detect_method(texts: Iterable[str]) -> str:
    joined = "\n".join(texts)
    methods = []
    for method in (
        "假设开发法",
        "上市公司比较法",
        "交易案例比较法",
        "市场法",
        "收益法",
        "重置成本法",
        "成本法",
        "资产基础法",
    ):
        if method in joined and method not in methods:
            methods.append(method)
    return "、".join(methods)


def workpaper_table_options(rules: Dict[str, Any]) -> Dict[str, Any]:
    options = dict(DEFAULT_WORKPAPER_TABLE_OPTIONS)
    configured = rules.get("workpaper_tables")
    if isinstance(configured, dict):
        options.update(configured)
    return options


def option_int(options: Dict[str, Any], key: str, default: int) -> int:
    try:
        value = int(options.get(key, default))
    except (TypeError, ValueError):
        return default
    return value if value > 0 else default


def keyword_hits(text: str, keywords: Sequence[str]) -> List[str]:
    source = text.upper()
    hits = []
    for keyword in keywords:
        if keyword.upper() in source and keyword not in hits:
            hits.append(keyword)
    return hits


def row_nonempty_texts(row: Sequence[Any]) -> List[str]:
    return [clean_text(cell) for cell in row if clean_text(cell)]


def split_nonempty_row_blocks(
    matrix: Sequence[Sequence[Any]],
) -> List[List[Tuple[int, Sequence[Any]]]]:
    blocks: List[List[Tuple[int, Sequence[Any]]]] = []
    current: List[Tuple[int, Sequence[Any]]] = []
    for index, row in enumerate(matrix):
        if row_nonempty_texts(row):
            current.append((index, row))
            continue
        if current:
            blocks.append(current)
            current = []
    if current:
        blocks.append(current)
    return blocks


def choose_workpaper_header_index(
    block: Sequence[Tuple[int, Sequence[Any]]],
) -> int:
    limit = min(len(block), 12)
    for index, (_source_index, row) in enumerate(block[:limit]):
        texts = row_nonempty_texts(row)
        if len(texts) >= 2 and keyword_hits(" ".join(texts), WORKPAPER_HEADER_KEYWORDS):
            return index
    for index, (_source_index, row) in enumerate(block[:limit]):
        if len(row_nonempty_texts(row)) >= 3:
            return index
    for index, (_source_index, row) in enumerate(block[:limit]):
        if len(row_nonempty_texts(row)) >= 2:
            return index
    return -1


def workpaper_table_title(
    path: Path,
    sheet_name: str,
    block: Sequence[Tuple[int, Sequence[Any]]],
    header_index: int,
) -> str:
    if header_index <= 0:
        return clean_text(sheet_name) or path.stem
    metadata_markers = ("评估基准日", "被评估单位", "金额单位", "单位：")
    for _source_index, row in block[: max(header_index, 1)]:
        texts = row_nonempty_texts(row)
        if not texts:
            continue
        title = " ".join(texts)
        if (
            len(texts) <= 2
            and len(title) <= 80
            and not any(marker in title for marker in metadata_markers)
        ):
            return title
    return clean_text(sheet_name) or path.stem


def infer_workpaper_subject_hints(text: str) -> List[str]:
    normalized = normalize_name(text)
    matches = []
    for subject, aliases in WORKPAPER_SUBJECT_HINTS:
        candidates = (subject, *aliases)
        if any(normalize_name(candidate) in normalized for candidate in candidates):
            matches.append(subject)
    return matches


def workpaper_table_excluded(search_text: str) -> bool:
    if any(keyword in search_text for keyword in WORKPAPER_EXCLUDE_TABLE_KEYWORDS):
        return True
    if "资产负债表" not in search_text:
        is_asset_detail = "资产编号" in search_text and "账面价值" in search_text
        keep_markers = (
            "比较因素",
            "案例",
            "实例",
            "参数",
            "价格指数",
            "调整系数表",
            "成新率计算表",
        )
        return is_asset_detail and not any(marker in search_text for marker in keep_markers)
    keep_markers = ("市场法评估明细表", "财务指标", "可比", "案例", "修正")
    return not any(marker in search_text for marker in keep_markers)


def build_workpaper_table(
    sheet: Any,
    path: Path,
    sheet_name: str,
    block: Sequence[Tuple[int, Sequence[Any]]],
    options: Dict[str, Any],
) -> Optional[Dict[str, Any]]:
    header_index = choose_workpaper_header_index(block)
    if header_index < 0 or header_index >= len(block) - 1:
        return None
    header_source_index, header_row = block[header_index]
    data_rows = [
        (source_index, row)
        for source_index, row in block[header_index + 1 :]
        if row_nonempty_texts(row)
    ]
    if not data_rows:
        return None
    raw_headers, raw_rows, used_columns = trim_table_columns(
        header_row,
        [row for _source_index, row in data_rows],
    )
    if len(raw_headers) < 2 or not raw_rows:
        return None

    max_rows = option_int(options, "max_rows_per_table", 80)
    max_columns = option_int(options, "max_columns_per_table", 40)
    if len(raw_rows) > max_rows or len(raw_headers) > max_columns:
        return None

    headers = []
    for index, header in enumerate(raw_headers):
        used_column = row_value(used_columns, index) or 0
        display = cell_display_text(sheet, header_source_index, used_column)
        headers.append(display or clean_text(header) or f"列{index + 1}")

    rows = []
    for source_index, row in data_rows:
        formatted = []
        for index, _header in enumerate(headers):
            used_column = row_value(used_columns, index) or 0
            display = cell_display_text(sheet, source_index, used_column)
            value = row_value(row, used_column)
            if display and display != "###":
                formatted.append(display)
            else:
                formatted.append(format_table_cell(value, row_value(headers, index)))
        if any(clean_text(cell) for cell in formatted):
            rows.append(formatted)
    if not rows:
        return None

    title = workpaper_table_title(path, sheet_name, block, header_index)
    title_scope = f"{sheet_name} {title}"
    if any(marker in title_scope for marker in ("资产表", "负债表", "资产负债表", "租赁合同台账")):
        return None
    if "评估底稿" in title_scope and not any(
        marker in title_scope
        for marker in ("案例", "实例", "参数", "指数", "修正系数", "调整系数", "成新率")
    ):
        return None
    preview_values = [
        path.name,
        str(path.parent),
        sheet_name,
        title,
        *headers,
        *(cell for row in rows[:12] for cell in row),
    ]
    search_text = " ".join(clean_text(value) for value in preview_values)
    hits = keyword_hits(
        search_text,
        tuple(options.get("keywords") or WORKPAPER_TABLE_KEYWORDS),
    )
    if not hits or workpaper_table_excluded(search_text):
        return None

    subject_text = " ".join(
        clean_text(value)
        for value in (
            sheet_name,
            title,
            *headers,
            *(cell for row in rows[:12] for cell in row),
        )
    )
    subject_hints = infer_workpaper_subject_hints(subject_text)
    if not subject_hints:
        subject_hints = infer_workpaper_subject_hints(search_text)
    title_scope_key = normalize_name(title_scope)
    if "构筑物" in title_scope_key and "房屋建筑物" not in title_scope_key:
        subject_hints = [
            hint for hint in subject_hints if hint != "房屋建筑物"
        ]
    if "房屋建筑物" in title_scope_key and "构筑物" not in title_scope_key:
        subject_hints = [
            hint for hint in subject_hints if hint != "构筑物及其他辅助设施"
        ]

    return {
        "title": title,
        "source_file": str(path),
        "file_name": path.name,
        "sheet": sheet_name,
        "headers": headers,
        "rows": rows,
        "row_count": len(rows),
        "column_count": len(headers),
        "keywords": hits,
        "subject_hints": subject_hints,
    }


def extract_workpaper_tables(
    sheets: Sequence[Any],
    path: Path,
    rules: Dict[str, Any],
) -> List[Dict[str, Any]]:
    options = workpaper_table_options(rules)
    if not bool(options.get("enabled", True)):
        return []
    max_scan_rows = option_int(options, "max_scan_rows", 500)
    max_scan_columns = option_int(options, "max_scan_columns", 40)
    max_tables = option_int(options, "max_tables_per_workpaper", 12)
    tables = []
    seen = set()
    for sheet in sheets:
        sheet_name = clean_text(sheet.Name)
        matrix = read_range(
            sheet,
            max_rows=max_scan_rows,
            max_columns=max_scan_columns,
        )
        for block in split_nonempty_row_blocks(matrix):
            table = build_workpaper_table(
                sheet,
                path,
                sheet_name,
                block,
                options,
            )
            if not table:
                continue
            key = (
                table["file_name"],
                table["sheet"],
                table["title"],
                tuple(table["headers"]),
            )
            if key in seen:
                continue
            seen.add(key)
            tables.append(table)
            if len(tables) >= max_tables:
                return tables
    return tables


def summarize_workpaper(
    excel: Any,
    path: Path,
    rules: Dict[str, Any],
) -> Dict[str, Any]:
    with open_workbook(excel, path) as workbook:
        sheets = visible_worksheets(workbook)
        key_text = []
        sheet_names = []
        for sheet in sheets:
            sheet_names.append(clean_text(sheet.Name))
            matrix = read_range(sheet, max_rows=15, max_columns=20)
            for row in matrix:
                for cell in row:
                    text = clean_text(cell)
                    if text and text not in key_text:
                        key_text.append(text)
                    if len(key_text) >= 80:
                        break
                if len(key_text) >= 80:
                    break
            if len(key_text) >= 80:
                break
        tables = extract_workpaper_tables(sheets, path, rules)
        return {
            "file": str(path),
            "kind": detect_method([path.name, *sheet_names, *key_text]),
            "visible_sheets": sheet_names,
            "valuation_method": detect_method(key_text),
            "key_text": key_text[:40],
            "tables": tables,
            "table_count": len(tables),
            "warnings": [],
        }


def build_input_data(
    input_dir: Path,
    excel_files: Sequence[Path],
    rules: Dict[str, Any],
) -> Dict[str, Any]:
    detail_paths = [path for path in excel_files if "评估明细表" in path.name]
    candidates = []
    candidate_errors = []
    with excel_application() as excel:
        for index, path in enumerate(detail_paths, start=1):
            print(f"[明细表 {index}/{len(detail_paths)}] {path.name}")
            try:
                candidates.append(read_detail_candidate(excel, path))
            except Exception as exc:
                candidate_errors.append(f"{path.name}: {exc}")
        primary = choose_primary_workbook(input_dir, candidates)
        print(f"母公司明细表: {Path(primary['file']).name}")
        print(f"被评估单位: {primary['entity']}")
        primary_details = summarize_primary_workbook(excel, primary, rules)

        related = [
            {
                "file": candidate["file"],
                "entity": candidate["entity"],
                "relationship": "subsidiary_or_related_entity",
                "used_for_parent_scope": False,
            }
            for candidate in candidates
            if candidate["file"] != primary["file"]
        ]
        related_candidates = [
            candidate
            for candidate in candidates
            if candidate["file"] != primary["file"]
        ]
        related_aliases = [
            alias
            for candidate in related_candidates
            for alias in company_aliases(candidate)
        ]
        all_workpaper_paths = [
            path
            for path in excel_files
            if path not in detail_paths
        ]
        workpaper_paths = [
            path
            for path in all_workpaper_paths
            if not any(alias in path.stem for alias in related_aliases)
        ]
        excluded_related_workpapers = [
            str(path)
            for path in all_workpaper_paths
            if path not in workpaper_paths
        ]
        workpapers = []
        for index, path in enumerate(workpaper_paths, start=1):
            print(f"[工作底稿 {index}/{len(workpaper_paths)}] {path.name}")
            try:
                workpapers.append(summarize_workpaper(excel, path, rules))
            except Exception as exc:
                workpapers.append(
                    {
                        "file": str(path),
                        "kind": "",
                        "visible_sheets": [],
                        "valuation_method": "",
                        "key_text": [],
                        "tables": [],
                        "table_count": 0,
                        "warnings": [str(exc)],
                    }
                )

    classification_rows = primary["classification"]["rows"]
    excluded_subjects = set(rules.get("summary_categories", []))
    excluded_subjects.update({"负债合计", "净资产", "所有者权益"})
    subject_order = [
        row["科目名称"]
        for row in classification_rows
        if (
            numeric_value(row.get("账面价值")) not in (None, 0)
            or numeric_value(row.get("评估价值")) not in (None, 0)
        )
        and row["科目名称"] not in excluded_subjects
    ]
    physical_subjects = set(rules.get("physical_asset_subjects", []))
    physical_assets = [
        {
            "subject": item["subject"],
            "item_count": item["item_count"],
            "main_names": item["content_summary"],
            "locations": item["location_summary"],
            "source_sheets": [item["sheet"]],
            "source_workpapers": [],
        }
        for item in primary_details["subjects"]
        if item["subject"] in physical_subjects
        or any(
            marker in item["subject"]
            for marker in (
                "存货",
                "原材料",
                "产成品",
                "在产品",
                "周转材料",
                "发出商品",
                "委托加工物资",
                "房屋",
                "构筑物",
                "设备",
                "车辆",
                "土地",
                "在建工程",
            )
        )
    ]
    return {
        "schema_version": "1.0",
        "document_type": "2",
        "source_type": "valuation_project_folder",
        "source_root": str(input_dir),
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "project": {
            "被评估单位": primary["entity"],
            "评估基准日": primary["benchmark_date"],
            "评估对象": f"{primary['entity']}于评估基准日的股东全部权益价值",
            "最终评估方法": "",
        },
        "primary_workbook": primary_details,
        "related_workbooks": related,
        "asset_basis_summary": primary["summary"],
        "classification_summary": primary["classification"],
        "subject_order": subject_order,
        "subjects": primary_details["subjects"],
        "physical_assets": physical_assets,
        "off_balance_assets": [],
        "workpapers": workpapers,
        "excluded_related_workpapers": excluded_related_workpapers,
        "warnings": candidate_errors,
        "errors": [],
    }


def select_method_library(
    library: Dict[str, Any],
    subject_order: Sequence[str],
) -> List[Dict[str, Any]]:
    aliases = library.get("aliases", {})
    entries = library.get("entries", [])
    selected = []
    used = set()
    for subject in subject_order:
        target = aliases.get(subject, subject)
        target_names = [normalize_name(target)]
        if any(marker in subject for marker in ("房屋", "构筑物")):
            target_names.append(normalize_name("房屋建（构）筑物"))
        if "技术类无形资产" in subject:
            target_names.extend(
                (
                    normalize_name("专利权"),
                    normalize_name("商标权"),
                )
            )
        matches = [
            entry
            for entry in entries
            if any(
                target_name in normalize_name(entry.get("subject"))
                or normalize_name(entry.get("subject")) in target_name
                for target_name in target_names
            )
        ]
        for entry in matches:
            key = entry.get("subject")
            if key in used:
                continue
            used.add(key)
            selected.append(
                {
                    "subject": key,
                    "plain_text": str(entry.get("plain_text") or "").strip()[:16000],
                }
            )
    return selected


def compact_workpapers_for_ai(workpapers: Any) -> List[Dict[str, Any]]:
    if not isinstance(workpapers, list):
        return []
    compact = []
    for item in workpapers:
        if not isinstance(item, dict):
            continue
        tables = []
        for table in item.get("tables") or []:
            if not isinstance(table, dict):
                continue
            tables.append(
                {
                    "title": table.get("title"),
                    "sheet": table.get("sheet"),
                    "file_name": table.get("file_name"),
                    "row_count": table.get("row_count"),
                    "column_count": table.get("column_count"),
                    "keywords": table.get("keywords"),
                    "subject_hints": table.get("subject_hints"),
                }
            )
        compact.append(
            {
                "file": item.get("file"),
                "kind": item.get("kind"),
                "visible_sheets": item.get("visible_sheets"),
                "valuation_method": item.get("valuation_method"),
                "key_text": item.get("key_text"),
                "table_count": item.get("table_count"),
                "tables": tables,
                "warnings": item.get("warnings"),
            }
        )
    return compact


def compact_ai_input(data: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "project": data["project"],
        "asset_basis_summary": data["asset_basis_summary"],
        "classification_summary": data["classification_summary"],
        "subject_order": data["subject_order"],
        "subjects": [
            {
                key: item.get(key)
                for key in (
                    "category",
                    "subject",
                    "sheet",
                    "unit",
                    "book_value",
                    "assessed_value",
                    "increase_decrease",
                    "increase_rate",
                    "item_count",
                    "content_summary",
                    "location_summary",
                    "main_age_bucket",
                    "representative_rows",
                    "warnings",
                )
            }
            for item in data["subjects"]
        ],
        "physical_assets": data["physical_assets"],
        "workpapers": compact_workpapers_for_ai(data.get("workpapers")),
        "method_library": data.get("method_library", []),
        "warnings": data["warnings"],
    }


def call_gemini(
    *,
    api_key: str,
    model: str,
    thinking_level: str,
    prompt_text: str,
    data: Dict[str, Any],
    store: bool,
) -> Tuple[Dict[str, Any], Dict[str, Any]]:
    runtime_prompt = (
        f"{prompt_text}\n\n"
        "## 本次输出结构补充要求\n"
        "只返回 JSON。scope_intro 只写范围表之前的段落；"
        "asset_descriptions 写范围表之后的委估资产负债、表外资产及引用报告段落；"
        "method_sections 和 technical_sections 按 subject_order 顺序输出。"
        "不要在段落中重复生成表格，汇总表由程序根据源 Excel 写入。"
        "最终评估方法不明确时，不得擅自选择收益法、市场法或资产基础法作为最终结论。\n\n"
        "## 本次结构化数据\n"
        + json.dumps(compact_ai_input(data), ensure_ascii=False, indent=2)
    )
    client = create_gemini_client(api_key)
    interaction = client.create(
        model=model,
        input=runtime_prompt,
        response_format={
            "type": "text",
            "mime_type": "application/json",
            "schema": RESPONSE_SCHEMA,
        },
        generation_config={"thinking_level": thinking_level},
        store=store,
    )
    if interaction.get("status") == "failed":
        raise RuntimeError(
            f"Gemini interaction 失败: {interaction.get('error') or interaction}"
        )
    output_text = get_interaction_output_text(interaction)
    if not output_text:
        raise RuntimeError("Gemini 返回了空响应")
    try:
        payload = json.loads(output_text)
    except json.JSONDecodeError as exc:
        raise RuntimeError(f"Gemini 返回的结构化 JSON 无法解析: {exc}") from exc
    metadata = {
        "interaction_id": clean_text(interaction.get("id")),
        "model": clean_text(interaction.get("model")) or model,
        "usage": interaction.get("usage")
        if isinstance(interaction.get("usage"), dict)
        else {},
    }
    return payload, metadata


def paragraph_before(anchor: Paragraph, text: str, style: Optional[str] = None) -> Paragraph:
    element = OxmlElement("w:p")
    anchor._p.addprevious(element)
    paragraph = Paragraph(element, anchor._parent)
    if style:
        try:
            paragraph.style = style
        except KeyError:
            pass
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    if style not in {"报告章标题", "报告节标题"}:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
    run = paragraph.add_run(clean_text(text))
    run.font.name = "宋体" if style not in {"报告章标题", "报告节标题"} else "黑体"
    run._element.rPr.rFonts.set(
        qn("w:eastAsia"),
        "宋体" if style not in {"报告章标题", "报告节标题"} else "黑体",
    )
    run.font.size = Pt(12 if style not in {"报告章标题", "报告节标题"} else 14)
    run.font.bold = style in {"报告章标题", "报告节标题"}
    return paragraph


def table_before(
    document: Document,
    anchor: Paragraph,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    *,
    font_size: int = 9,
    format_cells: bool = True,
) -> Table:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = clean_text(value)
    if format_cells:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = "宋体"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                        run.font.size = Pt(font_size)
    anchor._p.addprevious(table._tbl)
    return table


def detail_table_before(
    document: Document,
    anchor: Paragraph,
    detail_table: Dict[str, Any],
) -> Optional[Table]:
    headers = detail_table.get("headers") or []
    rows = detail_table.get("rows") or []
    if not headers or not rows:
        return None
    return table_before(
        document,
        anchor,
        headers,
        rows,
        font_size=7,
        format_cells=False,
    )


def find_placeholder(document: Document, placeholder: str) -> Paragraph:
    expected = "{{" + placeholder + "}}"
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == expected:
            return paragraph
    raise ValueError(f"模板中未找到占位符: {expected}")


def remove_paragraph(paragraph: Paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def replace_header_placeholder(document: Document, name: str, value: str) -> None:
    token = "{{" + name + "}}"
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            if token in paragraph.text:
                for run in paragraph.runs:
                    if token in run.text:
                        run.text = run.text.replace(token, value)


def scope_table_rows(summary_rows: Sequence[Dict[str, Any]]) -> List[List[str]]:
    fixed = {
        "流动资产",
        "非流动资产",
        "资产总计",
        "流动负债",
        "非流动负债",
        "负债总计",
        "股东全部权益价值",
        "股东全部权益价值（单体）",
        "股东全部权益价值（合并）",
        "净资产（所有者权益）",
    }
    rows = []
    for row in summary_rows:
        project = clean_text(row.get("项目"))
        book = numeric_value(row.get("账面价值"))
        if project not in fixed and book in (None, 0):
            continue
        rows.append(
            [
                clean_text(row.get("序号")),
                project,
                format_number(book),
            ]
        )
    return rows


def conclusion_table_rows(
    summary_rows: Sequence[Dict[str, Any]],
) -> List[List[str]]:
    rows = []
    for row in summary_rows:
        book = numeric_value(row.get("账面价值"))
        assessed = numeric_value(row.get("评估价值"))
        if book in (None, 0) and assessed in (None, 0):
            continue
        rows.append(
            [
                clean_text(row.get("项目")),
                clean_text(row.get("序号")),
                format_number(book),
                format_number(assessed),
                format_number(row.get("增减值")),
                format_rate(row.get("增值率")),
            ]
        )
    return rows


def add_paragraphs(anchor: Paragraph, paragraphs: Any) -> None:
    if not isinstance(paragraphs, list):
        return
    for text in paragraphs:
        if clean_text(text):
            paragraph_before(anchor, clean_text(text))


def add_section_objects(anchor: Paragraph, sections: Any) -> None:
    if not isinstance(sections, list):
        return
    for section in sections:
        if not isinstance(section, dict):
            continue
        subject = clean_text(section.get("subject"))
        if subject:
            paragraph_before(anchor, subject, "报告节标题")
        add_paragraphs(anchor, section.get("paragraphs"))


def detail_table_for_subject(
    data: Dict[str, Any],
    subject: str,
) -> Dict[str, Any]:
    subject_key = normalize_name(subject)
    for item in data.get("subjects", []):
        if normalize_name(item.get("subject")) == subject_key:
            table = item.get("detail_table")
            return table if isinstance(table, dict) else {}
    return {}


def add_detail_table_for_subject(
    document: Document,
    anchor: Paragraph,
    data: Dict[str, Any],
    subject: str,
    *,
    placement: str,
) -> None:
    options = data.get("_detail_table_options")
    if not isinstance(options, dict):
        options = {}
    if not bool(options.get("enabled", True)):
        return
    if placement == "report_scope" and not bool(
        options.get("insert_in_report_scope", True)
    ):
        return
    if placement == "explanation_technical" and not bool(
        options.get("insert_in_explanation_technical", True)
    ):
        return
    subject_key = normalize_name(subject)
    allowed_key = (
        "report_scope_subjects"
        if placement == "report_scope"
        else "explanation_technical_subjects"
    )
    allowed_subjects = options.get(allowed_key)
    if isinstance(allowed_subjects, list) and allowed_subjects:
        allowed = {normalize_name(item) for item in allowed_subjects}
        if subject_key not in allowed:
            return
    table = detail_table_for_subject(data, subject)
    if not table.get("rows"):
        return
    max_rows = int(options.get("max_rows_per_detail_table") or 0)
    if max_rows > 0 and len(table.get("rows") or []) > max_rows:
        return
    row_count = len(table.get("rows") or [])
    sheet_name = clean_text(table.get("sheet"))
    paragraph_before(
        anchor,
        f"{subject}明细表如下（来源：{sheet_name}，共{row_count}行）：",
    )
    detail_table_before(document, anchor, table)


def subject_matches_workpaper_table(subject: str, table: Dict[str, Any]) -> bool:
    subject_text = clean_subject_name(clean_text(subject))
    subject_key = normalize_name(subject_text)
    if not subject_key:
        return False
    title_scope_key = normalize_name(
        f"{clean_text(table.get('sheet'))} {clean_text(table.get('title'))}"
    )
    if (
        normalize_name("房屋建筑物") in subject_key
        and "构筑物" in title_scope_key
        and "房屋建筑物" not in title_scope_key
    ):
        return False
    if (
        normalize_name("构筑物及其他辅助设施") in subject_key
        and "房屋建筑物" in title_scope_key
        and "构筑物" not in title_scope_key
    ):
        return False
    if (
        normalize_name("构筑物及其他辅助设施") in subject_key
        and "租金" in title_scope_key
        and "构筑物" not in title_scope_key
    ):
        return False

    hint_keys = [normalize_name(hint) for hint in table.get("subject_hints") or []]
    if any(
        subject_key == hint_key or subject_key in hint_key or hint_key in subject_key
        for hint_key in hint_keys
    ):
        return True

    metadata = " ".join(
        clean_text(value)
        for value in (
            table.get("sheet"),
            table.get("title"),
            *(table.get("keywords") or []),
            *(table.get("headers") or []),
        )
    )
    metadata_key = normalize_name(metadata)
    if subject_key in metadata_key:
        return True

    for canonical, aliases in WORKPAPER_SUBJECT_HINTS:
        canonical_key = normalize_name(canonical)
        alias_keys = [normalize_name(alias) for alias in aliases]
        subject_is_alias = (
            subject_key == canonical_key
            or canonical_key in subject_key
            or any(alias_key and alias_key in subject_key for alias_key in alias_keys)
        )
        if not subject_is_alias:
            continue
        if canonical_key in hint_keys or any(
            alias_key and alias_key in metadata_key for alias_key in alias_keys
        ):
            return True
    return False


def workpaper_tables_for_subject(
    data: Dict[str, Any],
    subject: str,
) -> List[Dict[str, Any]]:
    if normalize_name(subject) in {
        normalize_name("固定资产"),
        normalize_name("无形资产"),
        normalize_name("流动资产"),
        normalize_name("非流动资产"),
    }:
        return []

    options = data.get("_workpaper_table_options")
    if not isinstance(options, dict):
        options = {}
    if not bool(options.get("enabled", True)):
        return []
    if not bool(options.get("insert_in_explanation_technical", True)):
        return []

    allowed_subjects = options.get("explanation_technical_subjects")
    if isinstance(allowed_subjects, list) and allowed_subjects:
        allowed = {normalize_name(item) for item in allowed_subjects}
        if normalize_name(subject) not in allowed:
            return []

    max_tables = option_int(options, "max_tables_per_subject", 10)
    matched = []
    seen = set()
    for workpaper in data.get("workpapers") or []:
        if not isinstance(workpaper, dict):
            continue
        for table in workpaper.get("tables") or []:
            if not isinstance(table, dict):
                continue
            if not table.get("headers") or not table.get("rows"):
                continue
            if not subject_matches_workpaper_table(subject, table):
                continue
            key = (
                table.get("sheet"),
                table.get("title"),
                tuple(table.get("headers") or []),
            )
            if key in seen:
                continue
            seen.add(key)
            matched.append(table)
            if len(matched) >= max_tables:
                return matched
    return matched


def add_workpaper_tables_for_subject(
    document: Document,
    anchor: Paragraph,
    data: Dict[str, Any],
    subject: str,
) -> None:
    tables = workpaper_tables_for_subject(data, subject)
    for table in tables:
        title = clean_text(table.get("title")) or clean_text(table.get("sheet"))
        file_name = clean_text(table.get("file_name")) or Path(
            clean_text(table.get("source_file"))
        ).name
        sheet_name = clean_text(table.get("sheet"))
        row_count = len(table.get("rows") or [])
        paragraph_before(
            anchor,
            f"专项底稿表：{title}（来源：{file_name} / {sheet_name}，共{row_count}行）",
        )
        detail_table_before(document, anchor, table)


def add_asset_descriptions(
    document: Document,
    anchor: Paragraph,
    data: Dict[str, Any],
    paragraphs: Any,
) -> None:
    if not isinstance(paragraphs, list):
        return
    subject_names = {
        normalize_name(item.get("subject")): clean_text(item.get("subject"))
        for item in data.get("subjects", [])
        if isinstance(item.get("detail_table"), dict)
        and item.get("detail_table", {}).get("rows")
    }
    pending_subject = ""
    for text in paragraphs:
        text = clean_text(text)
        if not text:
            continue
        paragraph_before(anchor, text)
        normalized = normalize_name(text)
        if normalized in subject_names:
            pending_subject = subject_names[normalized]
            continue
        if pending_subject:
            add_detail_table_for_subject(
                document,
                anchor,
                data,
                pending_subject,
                placement="report_scope",
            )
            pending_subject = ""
    if pending_subject:
        add_detail_table_for_subject(
            document,
            anchor,
            data,
            pending_subject,
            placement="report_scope",
        )


def add_section_objects_with_detail_tables(
    document: Document,
    anchor: Paragraph,
    data: Dict[str, Any],
    sections: Any,
) -> None:
    if not isinstance(sections, list):
        return
    for section in sections:
        if not isinstance(section, dict):
            continue
        subject = clean_text(section.get("subject"))
        if subject:
            paragraph_before(anchor, subject, "报告节标题")
        add_paragraphs(anchor, section.get("paragraphs"))
        if subject:
            add_detail_table_for_subject(
                document,
                anchor,
                data,
                subject,
                placement="explanation_technical",
            )
            add_workpaper_tables_for_subject(
                document,
                anchor,
                data,
                subject,
            )


def format_wanyuan_from_yuan(value: Any) -> str:
    number = numeric_value(value)
    if number is None:
        return "-"
    return format_number(number / 10000)


def find_summary_row(
    rows: Sequence[Dict[str, Any]],
    name: str,
) -> Dict[str, Any]:
    target = normalize_name(name)
    for row in rows:
        if normalize_name(row.get("项目")) == target:
            return row
    return {}


def subject_lookup(data: Dict[str, Any]) -> Dict[str, Dict[str, Any]]:
    return {
        normalize_name(item.get("subject")): item
        for item in data.get("subjects", [])
        if clean_text(item.get("subject"))
    }


def classification_lookup(data: Dict[str, Any]) -> Dict[str, Dict[str, Any]]:
    return {
        normalize_name(item.get("科目名称")): item
        for item in data.get("classification_summary", {}).get("rows", [])
        if clean_text(item.get("科目名称"))
    }


def clean_subject_name(subject: str) -> str:
    return clean_text(subject).replace("固定资产-", "").replace("无形资产-", "")


def nonzero_subject_order(data: Dict[str, Any]) -> List[str]:
    classes = classification_lookup(data)
    result = []
    for subject in data.get("subject_order", []):
        row = classes.get(normalize_name(subject), {})
        book = numeric_value(row.get("账面价值"))
        assessed = numeric_value(row.get("评估价值"))
        if book not in (None, 0) or assessed not in (None, 0):
            result.append(clean_subject_name(subject))
    return result


def valuation_subject_order(data: Dict[str, Any]) -> List[str]:
    ordered = []
    for subject in nonzero_subject_order(data):
        if subject not in ordered:
            ordered.append(subject)
    for item in data.get("subjects", []):
        subject = clean_subject_name(clean_text(item.get("subject")))
        if not subject:
            continue
        book = numeric_value(item.get("book_value"))
        assessed = numeric_value(item.get("assessed_value"))
        has_details = bool(item.get("item_count")) or bool(item.get("representative_rows"))
        if book in (None, 0) and assessed in (None, 0) and not has_details:
            continue
        if subject not in ordered:
            ordered.append(subject)
    return ordered


def subjects_in_category(
    data: Dict[str, Any],
    category: str,
) -> List[str]:
    classes = [
        row
        for row in data.get("classification_summary", {}).get("rows", [])
        if clean_text(row.get("类别")) == category
        and clean_text(row.get("科目名称"))
    ]
    ordered = []
    for row in classes:
        book = numeric_value(row.get("账面价值"))
        assessed = numeric_value(row.get("评估价值"))
        if book in (None, 0) and assessed in (None, 0):
            continue
        name = clean_subject_name(clean_text(row.get("科目名称")))
        if name not in ordered:
            ordered.append(name)
    return ordered


def join_cn(items: Sequence[str]) -> str:
    cleaned = [clean_text(item) for item in items if clean_text(item)]
    if not cleaned:
        return ""
    if len(cleaned) == 1:
        return cleaned[0]
    return "、".join(cleaned)


def top_values(
    values: Any,
    *,
    limit: int = 5,
) -> List[str]:
    if not isinstance(values, list):
        return []
    result = []
    for item in values:
        if isinstance(item, dict):
            value = clean_text(item.get("value"))
        else:
            value = clean_text(item)
        if value and value not in result:
            result.append(value)
        if len(result) >= limit:
            break
    return result


def describe_subject(
    subject: str,
    item: Dict[str, Any],
    class_row: Dict[str, Any],
) -> str:
    subject = clean_subject_name(subject)
    book = numeric_value(item.get("book_value"))
    if book is None:
        book = numeric_value(class_row.get("账面价值"))
    assessed = numeric_value(item.get("assessed_value"))
    if assessed is None:
        assessed = numeric_value(class_row.get("评估价值"))
    book_text = format_wanyuan_from_yuan(book)
    value_phrase = (
        f"账面价值{book_text}万元"
        if book is not None
        else "账面价值按评估明细表汇总列示"
    )
    item_count = item.get("item_count")
    count_text = f"，共{int(item_count)}项" if isinstance(item_count, int) else ""
    content = join_cn(top_values(item.get("content_summary"), limit=4))
    location = join_cn(top_values(item.get("location_summary"), limit=3))
    age = clean_text(item.get("main_age_bucket"))

    tail_parts = []
    if content:
        tail_parts.append(f"主要为{content}")
    if location:
        tail_parts.append(f"存放或坐落于{location}")
    if age:
        tail_parts.append(f"账龄以{age}为主")
    tail = "，" + "，".join(tail_parts) if tail_parts else ""

    if subject in {"应收账款", "其他应收款"}:
        return (
            f"{subject}{value_phrase}{count_text}{tail}。"
            "坏账准备按评估明细表列示口径处理。"
        )
    if subject in {"应付账款", "预收款项", "其他应付款"}:
        return f"{subject}{value_phrase}{count_text}{tail}。"
    if subject == "货币资金":
        return f"{subject}{value_phrase}{count_text}，包括现金、银行存款或其他货币资金等。"
    if assessed not in (None, 0) and book in (None, 0):
        return f"{subject}账面价值为零，评估价值{format_wanyuan_from_yuan(assessed)}万元。"
    return f"{subject}{value_phrase}{count_text}{tail}。"


def child_subject_names(parent: str) -> List[str]:
    mapping = {
        "存货": ["原材料", "产成品", "在产品", "库存商品", "周转材料", "发出商品"],
        "固定资产": [
            "房屋建筑物",
            "构筑物及其他辅助设施",
            "机器设备",
            "车辆",
            "电子设备",
        ],
        "无形资产": ["土地使用权", "其他无形资产", "技术类无形资产"],
    }
    return mapping.get(clean_subject_name(parent), [])


def build_programmatic_scope_sections(
    data: Dict[str, Any],
) -> Tuple[List[str], List[str]]:
    entity = clean_text(data["project"].get("被评估单位"))
    scope_intro = [
        "（一）评估对象",
        f"本次评估对象为{entity}于评估基准日的股东全部权益价值。",
        "（二）评估范围",
        (
            f"包括{entity}评估基准日未审资产负债表列示的全部资产和负债，"
            f"以及{entity}申报的表外资产，并且由{entity}提供的清单载明，"
            "具体评估范围请见本报告所附的“评估明细表”。"
        ),
        "企业申报的评估基准日表内资产及负债对应的会计报表未经审计，具体情况见下表：",
    ]

    summary_rows = data.get("asset_basis_summary", {}).get("rows", [])
    classes = classification_lookup(data)
    subjects = subject_lookup(data)
    category_names = []
    for name in ("流动资产", "非流动资产", "流动负债", "非流动负债"):
        row = find_summary_row(summary_rows, name)
        book = numeric_value(row.get("账面价值"))
        if book not in (None, 0):
            category_names.append(name)

    asset_descriptions = [
        "被评估单位填写的评估明细表内容除申报的表外资产外与未审资产负债表内容相一致。被评估单位已承诺无应纳入而未纳入本次评估范围的资产和负债。纳入评估范围的资产和负债与委托评估时确定的范围是一致的。",
        "委估主要资产及负债情况",
        (
            "纳入本次评估范围中的资产及负债包括"
            f"{join_cn(category_names)}，其中："
        ),
    ]
    for category in category_names:
        names = subjects_in_category(data, category)
        if not names:
            continue
        asset_descriptions.extend(
            [
                category,
                f"{category}包括{join_cn(names)}。",
            ]
        )
        for name in names:
            item = subjects.get(normalize_name(name), {})
            class_row = classes.get(normalize_name(name), {})
            if not item and not class_row:
                continue
            asset_descriptions.extend(
                [
                    name,
                    describe_subject(name, item, class_row),
                ]
            )
            for child_name in child_subject_names(name):
                child_item = subjects.get(normalize_name(child_name), {})
                if not child_item:
                    continue
                if not child_item.get("item_count") and not child_item.get(
                    "representative_rows"
                ):
                    continue
                asset_descriptions.extend(
                    [
                        child_name,
                        describe_subject(child_name, child_item, {}),
                    ]
                )
    off_balance = data.get("off_balance_assets") or []
    asset_descriptions.append("企业申报的表外资产情况")
    if off_balance:
        asset_descriptions.append(
            "企业申报的表外资产包括"
            + join_cn(clean_text(item.get("asset_type")) for item in off_balance)
            + "。"
        )
    else:
        asset_descriptions.append("本次评估企业无申报的表外资产。")
    asset_descriptions.extend(
        [
            "引用其他机构报告结论的情况",
            "本资产评估报告的评估结论无引用其他机构出具的报告的结论。",
        ]
    )
    return scope_intro, asset_descriptions


def build_programmatic_method_sections(data: Dict[str, Any]) -> List[Dict[str, Any]]:
    library = {
        normalize_name(item.get("subject")): str(item.get("plain_text") or "").strip()
        for item in data.get("method_library", [])
        if clean_text(item.get("subject"))
    }
    sections = []
    used = set()
    for subject in valuation_subject_order(data):
        candidates = [
            normalize_name(subject),
            normalize_name(subject.replace("固定资产", "")),
            normalize_name(subject.replace("无形资产", "")),
        ]
        text = ""
        for candidate in candidates:
            text = library.get(candidate, "")
            if text:
                break
        if not text:
            for key, value in library.items():
                if key and (key in normalize_name(subject) or normalize_name(subject) in key):
                    text = value
                    break
        if not text:
            continue
        paragraphs = [
            line.strip()
            for line in text.splitlines()
            if line.strip() and line.strip() != "参考模板："
        ]
        if paragraphs and normalize_name(paragraphs[0]).endswith(normalize_name(subject)):
            paragraphs = paragraphs[1:]
        if not paragraphs:
            continue
        key = normalize_name(subject)
        if key in used:
            continue
        used.add(key)
        sections.append({"subject": subject, "paragraphs": paragraphs})
    return sections


def format_cn_date(date_text: str) -> str:
    value = clean_text(date_text)
    match = re.fullmatch(r"(\d{4})[-/.年](\d{1,2})[-/.月](\d{1,2})日?", value)
    if not match:
        return value
    year, month, day = match.groups()
    return f"{year}年{int(month)}月{int(day)}日"


def next_year_minus_one_day(date_text: str) -> str:
    value = clean_text(date_text)
    for pattern in ("%Y-%m-%d", "%Y.%m.%d", "%Y/%m/%d"):
        try:
            start = datetime.strptime(value, pattern).date()
            break
        except ValueError:
            start = None
    if start is None:
        return ""
    try:
        end = start.replace(year=start.year + 1) - timedelta(days=1)
    except ValueError:
        end = start.replace(year=start.year + 1, day=28) - timedelta(days=1)
    return f"{end.year}年{end.month}月{end.day}日"


def build_programmatic_conclusion_sections(
    data: Dict[str, Any],
) -> Tuple[List[str], List[str]]:
    entity = clean_text(data["project"].get("被评估单位"))
    benchmark = clean_text(data["project"].get("评估基准日"))
    rows = data.get("asset_basis_summary", {}).get("rows", [])
    assets = find_summary_row(rows, "资产总计")
    liabilities = find_summary_row(rows, "负债总计")
    equity = find_summary_row(rows, "股东全部权益价值") or find_summary_row(
        rows,
        "净资产",
    )

    conclusion_intro = [
        (
            "根据国家有关资产评估的法律、行政法规及资产评估准则的规定，"
            f"本着独立、客观、公正的原则及必要的评估程序，对{entity}"
            "的股东全部权益价值进行了评估。根据以上评估工作，得出如下评估结论："
        ),
        "资产基础法评估结果",
    ]
    if assets and liabilities and equity:
        conclusion_intro.append(
            (
                f"经资产基础法评估，{entity}总资产账面值"
                f"{format_number(assets.get('账面价值'))}万元，评估值"
                f"{format_number(assets.get('评估价值'))}万元，增值额"
                f"{format_number(assets.get('增减值'))}万元，增值率"
                f"{format_rate(assets.get('增值率'))}；负债账面值"
                f"{format_number(liabilities.get('账面价值'))}万元，评估值"
                f"{format_number(liabilities.get('评估价值'))}万元，增值额"
                f"{format_number(liabilities.get('增减值'))}万元，增值率"
                f"{format_rate(liabilities.get('增值率'))}；股东全部权益账面值"
                f"{format_number(equity.get('账面价值'))}万元，评估值"
                f"{format_number(equity.get('评估价值'))}万元，增值额"
                f"{format_number(equity.get('增减值'))}万元，增值率"
                f"{format_rate(equity.get('增值率'))}。资产基础法评估结果详见下表："
            )
        )
    conclusion_intro.append("资产评估结果汇总表（资产基础法）")

    final_conclusion = ["最终评估结论"]
    if equity:
        final_conclusion.append(
            (
                f"经评估，{entity}股东全部权益评估价值为"
                f"{format_number(equity.get('评估价值'))}万元，比账面价值"
                f"{format_number(equity.get('账面价值'))}万元增值"
                f"{format_number(equity.get('增减值'))}万元，增值率"
                f"{format_rate(equity.get('增值率'))}。"
            )
        )
    valid_to = next_year_minus_one_day(benchmark)
    final_conclusion.append("评估结论的使用有效期")
    if valid_to:
        final_conclusion.append(
            (
                "本评估报告所揭示的评估结论仅对评估报告中描述的经济行为有效，"
                "评估结论使用有效期为自评估基准日起一年，即自评估基准日"
                f"{format_cn_date(benchmark)}起至{valid_to}止。超过一年使用本资产评估报告所列示的评估结论无效。"
                "国家法律、行政法规另有规定的除外。"
            )
        )
    return conclusion_intro, final_conclusion


def build_programmatic_technical_sections(data: Dict[str, Any]) -> List[Dict[str, Any]]:
    classes = classification_lookup(data)
    subjects = subject_lookup(data)
    method_sections = {
        normalize_name(section.get("subject")): section.get("paragraphs", [])
        for section in build_programmatic_method_sections(data)
    }
    sections = []
    for subject in valuation_subject_order(data):
        item = subjects.get(normalize_name(subject), {})
        class_row = classes.get(normalize_name(subject), {})
        book = numeric_value(item.get("book_value"))
        if book is None:
            book = numeric_value(class_row.get("账面价值"))
        assessed = numeric_value(item.get("assessed_value"))
        if assessed is None:
            assessed = numeric_value(class_row.get("评估价值"))
        paragraphs = [
            "评估范围",
            f"本次评估范围为{subject}，账面价值{format_wanyuan_from_yuan(book)}万元。",
            "资产概况",
            describe_subject(subject, item, class_row),
            "评估过程",
            "准备阶段",
            "对确定的评估范围内的资产或负债的构成情况进行初步了解，提交评估准备资料清单和评估申报明细表标准格式，指导企业填写评估申报明细表。",
            "现场调查阶段",
            "根据企业提供的评估申报资料，进行总账、明细账、会计报表及评估明细表的核对，使之相符；对内容不符、重复申报、遗漏未报项目进行改正，由企业重新填报。",
            "评定估算阶段",
            "在账务核对清晰、情况了解清楚并已收集到评估所需资料的基础上进行评定估算，编制评估明细表和汇总表，撰写评估技术说明。",
            "评估方法",
        ]
        paragraphs.extend(method_sections.get(normalize_name(subject), []))
        paragraphs.extend(
            [
                "评估结果",
                f"经评估，{subject}评估值为{format_wanyuan_from_yuan(assessed)}万元。",
            ]
        )
        sections.append({"subject": subject, "paragraphs": paragraphs})
    return sections


def enrich_generated_sections(
    data: Dict[str, Any],
    generated: Dict[str, Any],
    rules: Dict[str, Any],
) -> Dict[str, Any]:
    generated = json.loads(json.dumps(generated, ensure_ascii=False))
    data["_detail_table_options"] = rules.get("detail_tables", {})
    data["_workpaper_table_options"] = workpaper_table_options(rules)
    report = generated.setdefault("report", {})
    explanation = generated.setdefault("explanation", {})
    scope_intro, asset_descriptions = build_programmatic_scope_sections(data)
    conclusion_intro, final_conclusion = build_programmatic_conclusion_sections(data)
    method_sections = build_programmatic_method_sections(data)
    technical_sections = build_programmatic_technical_sections(data)

    fallbacks = rules.get("programmatic_section_fallbacks", {})
    if fallbacks.get("scope", True):
        report["scope_intro"] = scope_intro
        report["asset_descriptions"] = asset_descriptions
        explanation["scope_intro"] = scope_intro
    if fallbacks.get("method_sections", True) and method_sections:
        report["method_sections"] = method_sections
    if fallbacks.get("conclusion", True):
        report["conclusion_intro"] = conclusion_intro
        report["final_conclusion"] = final_conclusion
    if fallbacks.get("physical_assets", True):
        explanation["physical_assets"] = [
            "实物资产的分布情况及特点",
            *asset_descriptions,
        ]
    if fallbacks.get("technical_sections", True) and technical_sections:
        explanation["technical_sections"] = technical_sections
    return generated


def render_report(
    template_file: Path,
    output_file: Path,
    data: Dict[str, Any],
    generated: Dict[str, Any],
) -> None:
    document = Document(template_file)
    summary_rows = data["asset_basis_summary"]["rows"]
    report = generated["report"]

    scope_anchor = find_placeholder(document, "报告_评估对象和评估范围")
    add_paragraphs(scope_anchor, report.get("scope_intro"))
    paragraph_before(scope_anchor, "金额单位：人民币万元")
    table_before(
        document,
        scope_anchor,
        ("序号", "项目", "账面价值"),
        scope_table_rows(summary_rows),
    )
    add_asset_descriptions(
        document,
        scope_anchor,
        data,
        report.get("asset_descriptions"),
    )
    remove_paragraph(scope_anchor)

    method_anchor = find_placeholder(
        document,
        "报告_资产基础法具体评估方法介绍",
    )
    add_section_objects(method_anchor, report.get("method_sections"))
    remove_paragraph(method_anchor)

    conclusion_anchor = find_placeholder(document, "报告_评估结论")
    add_paragraphs(conclusion_anchor, report.get("conclusion_intro"))
    paragraph_before(conclusion_anchor, "金额单位：人民币万元")
    table_before(
        document,
        conclusion_anchor,
        ("项目", "序号", "账面价值", "评估价值", "增减值", "增值率"),
        conclusion_table_rows(summary_rows),
    )
    add_paragraphs(conclusion_anchor, report.get("final_conclusion"))
    remove_paragraph(conclusion_anchor)

    output_file.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_file)


def render_explanation(
    template_file: Path,
    output_file: Path,
    data: Dict[str, Any],
    generated: Dict[str, Any],
) -> None:
    document = Document(template_file)
    explanation = generated["explanation"]
    summary_rows = data["asset_basis_summary"]["rows"]
    entity = clean_text(data["project"]["被评估单位"])
    replace_header_placeholder(document, "项目名称", entity)

    scope_anchor = find_placeholder(document, "说明_评估对象与评估范围说明")
    add_paragraphs(scope_anchor, explanation.get("scope_intro"))
    paragraph_before(scope_anchor, "金额单位：人民币万元")
    table_before(
        document,
        scope_anchor,
        ("序号", "项目", "账面价值"),
        scope_table_rows(summary_rows),
    )
    remove_paragraph(scope_anchor)

    physical_anchor = find_placeholder(
        document,
        "说明_实物资产分布情况及特点",
    )
    add_paragraphs(physical_anchor, explanation.get("physical_assets"))
    remove_paragraph(physical_anchor)

    technical_anchor = find_placeholder(
        document,
        "说明_资产基础法评估技术说明",
    )
    add_section_objects_with_detail_tables(
        document,
        technical_anchor,
        data,
        explanation.get("technical_sections"),
    )
    remove_paragraph(technical_anchor)

    output_file.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_file)


def build_parser(project_root: Path) -> argparse.ArgumentParser:
    resource_root = (
        Path(getattr(sys, "_MEIPASS", "")) / "resources" / "2"
        if bool(getattr(sys, "frozen", False))
        else project_root / "bin" / "template" / "2"
    )
    parser = argparse.ArgumentParser(
        description="Generate type 2 enterprise valuation report documents."
    )
    parser.add_argument("--input-dir", required=True)
    parser.add_argument("--word-dir", default=str(project_root / "word"))
    parser.add_argument(
        "--records-file",
        default=str(project_root / "bin" / "json" / "类型2输入数据.json"),
    )
    parser.add_argument(
        "--report-template",
        default=str(resource_root / "企业价值评估报告-自动生成基底模板.docx"),
    )
    parser.add_argument(
        "--explanation-template",
        default=str(resource_root / "企业价值评估说明-自动生成基底模板.docx"),
    )
    parser.add_argument(
        "--prompt-file",
        default=str(resource_root / "企业价值评估报告自动生成-Prompt.md"),
    )
    parser.add_argument(
        "--rules-file",
        default=str(resource_root / "企业价值评估报告生成规则.json"),
    )
    parser.add_argument(
        "--method-library",
        default=str(resource_root / "资产基础法评估方法库.json"),
    )
    parser.add_argument("--model", default=None)
    parser.add_argument(
        "--api-key-file",
        default=str(project_root / "gemini_api.txt"),
    )
    parser.add_argument("--extract-only", action="store_true")
    parser.add_argument("--no-word-field-update", action="store_true")
    return parser


def get_project_root() -> Path:
    if not bool(getattr(sys, "frozen", False)):
        return Path(__file__).resolve().parent.parent
    executable_dir = Path(sys.executable).resolve().parent
    if (
        executable_dir.name.lower() == "dist"
        and (executable_dir.parent / "bin").is_dir()
    ):
        return executable_dir.parent
    return executable_dir


def main(argv: Optional[Sequence[str]] = None) -> int:
    project_root = get_project_root()
    args = build_parser(project_root).parse_args(argv)
    input_dir = Path(args.input_dir).expanduser().resolve()
    word_dir = Path(args.word_dir).expanduser().resolve()
    records_file = Path(args.records_file).expanduser().resolve()
    report_template = Path(args.report_template).expanduser().resolve()
    explanation_template = Path(args.explanation_template).expanduser().resolve()
    prompt_file = Path(args.prompt_file).expanduser().resolve()
    rules_file = Path(args.rules_file).expanduser().resolve()
    method_library_file = Path(args.method_library).expanduser().resolve()
    api_key_file = Path(args.api_key_file).expanduser().resolve()

    if not input_dir.is_dir():
        raise FileNotFoundError(f"项目目录不存在: {input_dir}")
    for label, path in (
        ("评估报告模板", report_template),
        ("评估说明模板", explanation_template),
        ("Prompt", prompt_file),
        ("生成规则", rules_file),
        ("评估方法库", method_library_file),
    ):
        if not path.is_file():
            raise FileNotFoundError(f"{label}不存在: {path}")

    rules = json.loads(rules_file.read_text(encoding="utf-8-sig"))
    prompt_text = prompt_file.read_text(encoding="utf-8-sig").strip()
    method_library = json.loads(
        method_library_file.read_text(encoding="utf-8-sig")
    )
    excel_files = discover_excel_files(input_dir)
    print(f"Excel files found: {len(excel_files)}")
    data = build_input_data(input_dir, excel_files, rules)
    method_subjects = [
        *data["subject_order"],
        *(item["subject"] for item in data["subjects"]),
    ]
    data["method_library"] = select_method_library(
        method_library,
        method_subjects,
    )
    records_file.parent.mkdir(parents=True, exist_ok=True)
    records_file.write_text(
        json.dumps(data, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    print(f"类型2 JSON: {records_file}")
    if args.extract_only:
        print("Extract-only mode: Word generation skipped.")
        return 0

    ai_config = rules.get("ai", {})
    model = clean_text(args.model) or clean_text(ai_config.get("model")) or DEFAULT_MODEL
    thinking_level = clean_text(ai_config.get("thinking_level")) or "medium"
    store = bool(ai_config.get("store_interactions", False))
    api_key = load_gemini_api_key(api_key_file)
    print(f"Calling Gemini: {model}")
    generated, interaction = call_gemini(
        api_key=api_key,
        model=model,
        thinking_level=thinking_level,
        prompt_text=prompt_text,
        data=data,
        store=store,
    )
    data["gemini_generated_sections"] = generated
    generated = enrich_generated_sections(data, generated, rules)
    data["generated_sections"] = generated
    data["gemini"] = interaction
    records_file.write_text(
        json.dumps(data, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )

    entity = sanitize_filename_part(data["project"]["被评估单位"])
    benchmark = sanitize_filename_part(data["project"]["评估基准日"])
    report_file = word_dir / f"1-评估报告-{entity}-{benchmark}.docx"
    explanation_file = word_dir / f"2-评估说明-{entity}-{benchmark}.docx"
    render_report(report_template, report_file, data, generated)
    render_explanation(explanation_template, explanation_file, data, generated)

    field_warnings = []
    if not args.no_word_field_update:
        for output_file in (report_file, explanation_file):
            warning = update_word_fields(output_file)
            if warning:
                field_warnings.append(f"{output_file.name}: {warning}")

    manifest = {
        "document_type": "2",
        "source_type": "valuation_project_folder",
        "source_root": str(input_dir),
        "records_file": str(records_file),
        "reports": [
            {"kind": "report", "filename": report_file.name},
            {"kind": "explanation", "filename": explanation_file.name},
        ],
        "gemini": interaction,
        "warnings": [*data.get("warnings", []), *generated.get("warnings", []), *field_warnings],
        "errors": [],
    }
    word_dir.mkdir(parents=True, exist_ok=True)
    manifest_file = word_dir / "generation_manifest.json"
    manifest_file.write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    print(f"Saved: {report_file}")
    print(f"Saved: {explanation_file}")
    print(f"Manifest: {manifest_file}")
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        raise SystemExit(1)
